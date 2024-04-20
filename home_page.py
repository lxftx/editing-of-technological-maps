import datetime
import os
import glob
import re

from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import QTimer
from PyQt5.QtWidgets import QWidget, QAbstractItemView
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from alert import AlertMessage
from doc_page import DocPage
from docx import Document


class HomePage(QWidget, AlertMessage):
    def __init__(self, main, search_widget, dir):
        super().__init__()
        self.home_page = self
        self.main = main
        self.row = None
        self.dir = dir
        self.timer = QTimer()

        self.search_widget = search_widget
        self.search_widget.clear()
        self.search_widget.textChanged.connect(self.search_users)

        self.home_content = QtWidgets.QLabel(self.home_page)
        self.home_content.setGeometry(QtCore.QRect(10, 10, 1021, 861))
        self.home_content.setStyleSheet("background-color:rgb(240, 240, 240);\n"
                                        "border-radius: 30px;\n"
                                        "color: rgb(255, 255, 17);")

        self.table_docx = QtWidgets.QTableWidget(self.home_page)
        self.table_docx.setEnabled(True)
        self.table_docx.setGeometry(QtCore.QRect(20, 50, 1001, 751))
        self.table_docx.setTabletTracking(False)
        self.table_docx.setObjectName("listWidget")
        self.table_docx.setColumnCount(5)  # Увеличиваем количество столбцов на 1
        self.table_docx.setRowCount(0)
        self.table_docx.setHorizontalHeaderLabels(
            ["Название", "Дата создания", "Дата обновления", "Кем создан", "Кем изменен"])
        self.table_docx.setEditTriggers(QAbstractItemView.NoEditTriggers)
        header = self.table_docx.horizontalHeader()
        header.setSectionResizeMode(QtWidgets.QHeaderView.Stretch)

        self.window_add_document = QtWidgets.QLabel(self.home_page)
        self.window_add_document.setGeometry(QtCore.QRect(250, 310, 561, 151))
        self.window_add_document.setStyleSheet("border:1px solid black; border-radius: 30px;")
        self.window_add_document.setText("")
        self.window_add_document.setObjectName("window_add_document")
        self.window_add_document.setVisible(False)

        self.window_add_document_title = QtWidgets.QLabel(self.home_page)
        self.window_add_document_title.setText('Имя документа:')
        self.window_add_document_title.setGeometry(QtCore.QRect(270, 320, 521, 41))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(16)
        self.window_add_document_title.setFont(font)
        self.window_add_document_title.setObjectName("window_add_document_title")
        self.window_add_document_title.setVisible(False)

        self.window_add_document_name_document = QtWidgets.QLineEdit(self.home_page)
        self.window_add_document_name_document.setGeometry(QtCore.QRect(270, 369, 521, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.window_add_document_name_document.setFont(font)
        self.window_add_document_name_document.setObjectName("window_add_document_name_document")
        self.window_add_document_name_document.setVisible(False)

        self.button_create = QtWidgets.QPushButton(self.home_page)
        self.button_create.setText('Создать')
        self.button_create.setGeometry(QtCore.QRect(270, 410, 241, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.button_create.setFont(font)
        self.button_create.setStyleSheet("QPushButton#button_create {\n"
                                         "    background-color: black; \n"
                                         "    color: white;\n"
                                         "    border-radius: 10px;\n"
                                         "}\n"
                                         "\n"
                                         "QPushButton#button_create:hover {\n"
                                         "    background-color: green;\n"
                                         "}")
        icon9 = QtGui.QIcon()
        icon9.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "yes.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.button_create.setIcon(icon9)
        self.button_create.setObjectName("button_create")
        self.button_create.setVisible(False)
        self.button_create.clicked.connect(self.create_document)

        self.button_cancel = QtWidgets.QPushButton(self.home_page)
        self.button_cancel.setText('Отменить')
        self.button_cancel.setGeometry(QtCore.QRect(550, 410, 241, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.button_cancel.setFont(font)
        self.button_cancel.setStyleSheet("QPushButton#button_cancel {\n"
                                         "    background-color: black; \n"
                                         "    color: white;\n"
                                         "    border-radius: 10px;\n"
                                         "}\n"
                                         "\n"
                                         "QPushButton#button_cancel:hover {\n"
                                         "    background-color: red;\n"
                                         "}")
        icon10 = QtGui.QIcon()
        icon10.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "cancel.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.button_cancel.setIcon(icon10)
        self.button_cancel.setIconSize(QtCore.QSize(16, 16))
        self.button_cancel.setObjectName("button_cancel")
        self.button_cancel.setVisible(False)
        self.button_cancel.clicked.connect(self.cancel_added_document)

        self.add_document = QtWidgets.QPushButton(self.home_page)
        self.add_document.setText('Создать тех.карту')
        self.add_document.setGeometry(QtCore.QRect(40, 820, 251, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.add_document.setFont(font)
        self.add_document.setStyleSheet("QPushButton#add_document {\n"
                                        "    background-color: black; \n"
                                        "    color: white;\n"
                                        "    border-radius: 10px;\n"
                                        "}\n"
                                        "\n"
                                        "QPushButton#add_document:hover {\n"
                                        "    background-color: green;\n"
                                        "}")
        icon7 = QtGui.QIcon()
        icon7.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "add_document.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.add_document.setIcon(icon7)
        self.add_document.setObjectName("add_document")
        self.add_document.clicked.connect(self.add_file)

        self.delete_document = QtWidgets.QPushButton(self.home_page)
        self.delete_document.setText('Удалить тех.карту')
        self.delete_document.setGeometry(QtCore.QRect(750, 820, 251, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.delete_document.setFont(font)
        self.delete_document.setStyleSheet("QPushButton#delete_document {\n"
                                           "    background-color: black; \n"
                                           "    color: white;\n"
                                           "    border-radius: 10px;\n"
                                           "}\n"
                                           "\n"
                                           "QPushButton#delete_document:hover {\n"
                                           "    background-color: red;\n"
                                           "}")
        icon8 = QtGui.QIcon()
        icon8.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "delete.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.delete_document.setIcon(icon8)
        self.delete_document.setObjectName("delete_document")
        self.delete_document.clicked.connect(self.del_document)

        self.alert = QtWidgets.QLabel(self.home_page)
        self.alert.setGeometry(QtCore.QRect(40, 270, 821, 201))
        self.alert.setStyleSheet("border-radius: 30px;")
        self.alert.setObjectName("alert_info")
        self.alert.setVisible(False)

        self.alert_title = QtWidgets.QLabel(self.home_page)
        self.alert_title.setGeometry(QtCore.QRect(40, 270, 961, 61))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(20)
        font.setBold(True)
        font.setWeight(75)
        self.alert_title.setFont(font)
        self.alert_title.setStyleSheet("border-top-left-radius: 30px;\n"
                                       "border-top-right-radius: 30px;\n"
                                       "background-color: rgb(227, 227, 227); \n"
                                       "color: rgb(179, 179, 179);")
        self.alert_title.setAlignment(QtCore.Qt.AlignCenter)
        self.alert_title.setObjectName("alert_info_title")
        self.alert_title.setText("Предупреждение")
        self.alert_title.setVisible(False)

        self.alert_text = QtWidgets.QLabel(self.home_page)
        self.alert_text.setGeometry(QtCore.QRect(40, 330, 961, 141))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(15)
        font.setUnderline(True)
        self.alert_text.setFont(font)
        self.alert_text.setStyleSheet("border-bottom-left-radius: 30px;\n"
                                      "border-bottom-right-radius: 30px;\n"
                                      "color: rgb(162, 162, 162);")
        self.alert_text.setAlignment(QtCore.Qt.AlignCenter)
        self.alert_text.setObjectName("alert_info_text")
        self.alert_text.setVisible(False)

        if self.dir:
            self.load_table()
            if hasattr(self, 'alert'):
                self.hide_alert()
        else:
            self.show_alert()
            self.alert_text.setText("Укажите путь к тех.картам через настройки")
        self.table_docx.cellClicked.connect(self.turn_document)
        self.table_docx.cellDoubleClicked.connect(self.open_document)

        if self.main.user.posts != "Технолог":
            css_button = "QPushButton {\n" \
                         "    background-color: rgb(115, 115, 115); \n" \
                         "    color: white;\n" \
                         "    border-radius: 10px;\n" \
                         "}\n"
            self.add_document.setEnabled(False)
            self.delete_document.setEnabled(False)
            self.add_document.setStyleSheet(css_button)
            self.delete_document.setStyleSheet(css_button)

        if self.main.burger_button.isChecked():
            self.open_sidebar()

    def get_path(self, dir, name):
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), dir, name)


    def closeEvent(self, event):
        self.main.db.disconnection_database()
        event.accept()

    def enable_widget(self, evt):
        self.table_docx.setEnabled(evt)
        self.add_document.setEnabled(evt)
        self.delete_document.setEnabled(evt)
        self.main.burger_button.setEnabled(evt)
        self.main.home_button.setEnabled(evt)
        self.main.users_button.setEnabled(evt)
        self.main.search_button.setEnabled(evt)
        self.main.user_info_button.setEnabled(evt)
        self.main.search_edit.setEnabled(evt)
        self.main.add_user_button.setEnabled(evt)
        self.main.exit_user_button.setEnabled(evt)
        self.window_add_document.setVisible(not evt)
        self.window_add_document_title.setVisible(not evt)
        self.window_add_document_name_document.setVisible(not evt)
        self.button_create.setVisible(not evt)
        self.button_cancel.setVisible(not evt)

    def search_users(self):
        self.table_docx.setRowCount(0)
        table_lst = []
        directory = self.dir
        text = self.search_widget.text().strip().title()
        files = os.listdir(directory)
        file_list = [directory + file for file in files if file.endswith('.docx')]
        for i, name in enumerate(file_list):
            if re.compile(text).search(name):
                # Получаем информацию о файле
                if not re.search(r"~\$", name):
                    file_name = os.path.basename(name)
                    creation_time = datetime.datetime.fromtimestamp(os.path.getctime(name)).strftime(
                        '%d.%m.%Y %H:%M:%S')
                    modification_time = datetime.datetime.fromtimestamp(os.path.getmtime(name)).strftime(
                        '%d.%m.%Y %H:%M:%S')
                    doc = Document(name)
                    core_props = doc.core_properties
                    authors = core_props.author
                    last_modified_by = core_props.last_modified_by
                    info_text = [file_name.split('.')[0], creation_time, modification_time, authors, last_modified_by]
                    table_lst.append(info_text)
        # Заполняем таблицу данными из списка
        self.table_docx.setRowCount(len(table_lst))
        for row_index, row_data in enumerate(table_lst):
            for col_index, col_data in enumerate(row_data):
                item = QtWidgets.QTableWidgetItem(col_data)
                self.table_docx.setItem(row_index, col_index, item)

    def create_document(self):
        name = self.window_add_document_name_document.text().strip()
        if name:
            if name not in [self.table_docx.item(i, 0).text() for i in range(self.table_docx.rowCount())]:
                # Создаем новый документ
                doc = Document(self.get_path(r'template_docx', "1.docx"))

                # Добавляем таблицу с двумя строками и двенадцатью столбцами
                table = doc.add_table(rows=3, cols=12)

                # Применяем стиль "Table Grid" к таблице
                table.style = "Table Grid"

                # Объединяем ячейки в первой строке
                for i in [[0, '№п/п'], [1, 'Наименование операций'], [2, 'Оборудование'], [3, 'Диаметр, мм'],
                          [5, '№ маршрута'], [6, 'Степень деформации, %'], [7, 'Режим  термообработки'],
                          [11, 'Примечание']]:
                    cell = table.cell(0, i[0])
                    if i[0] == 3:
                        next_cell = table.cell(0, i[0] + 1)
                        next_cell.text = i[1]
                        cell.merge(next_cell)  # Объединяем текущую ячейку с следующей
                    elif i[0] == 7:
                        next_cell = table.cell(0, i[0] + 3)
                        next_cell.text = i[1]
                        cell.merge(next_cell)  # Объединяем текущую ячейку с следующей
                    else:
                        cell.text = i[1]

                it = ['', '', '', 'До волочения', 'После волочения', '', '', 'Среда нагрева', 'Температура отжига, ̊С',
                      'Время, мин', 'Условия охлаждения']
                # Объединяем ячейки второй строки для указанных столбцов
                for i in range(12):
                    if i in (0, 1, 2, 5, 6, 11):
                        cell = table.cell(1, i)
                        next_cell = table.cell(0, i)
                        cell.merge(next_cell)  # Объединяем текущую ячейку соответствующей ячейкой из первой строки
                    else:
                        cell = table.cell(1, i)
                        cell.text = it[i]
                    cell = table.cell(2, i)
                    cell.text = str(i + 1)
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        run = paragraph.runs[0]
                        run.bold = True

                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(10)  # Используем Pt для установки размера шрифта
                        # Выравниваем текст по центру ячеек
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Сохраняем документ

                doc.save(self.dir+f"{name}.docx")
                self.enable_widget(True)
                file_name = name
                file_path = self.dir + file_name + ".docx"
                self.doc_page = DocPage(self.main, file_path)
                self.doc_page.setObjectName("doc_page")
                self.main.stackedWidget.addWidget(self.doc_page)
                self.main.stackedWidget.setCurrentWidget(self.doc_page)

            else:
                self.show_alert()
                self.alert_text.setText("Тех.карта с таким именем уже существует")
                self.timer.setSingleShot(True)
                self.timer.timeout.connect(self.hide_alert)
                self.timer.start(5000)
        else:
            self.show_alert()
            self.alert_text.setText("Поле имя документа обязательно к заполнению")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)

    def del_document(self):
        if self.row:
            file_path = self.dir+f'{self.row}.docx'
            try:
                os.remove(file_path)
                self.row = None
                self.load_table()
            except PermissionError:
                self.show_alert()
                self.alert_text.setText(
                    "Ошибка удаления! Файл был открыт, зайкроте его и снова выполните удаление документа!")
                self.timer.setSingleShot(True)
                self.timer.timeout.connect(self.hide_alert)
                self.timer.start(5000)

        else:
            self.show_alert()
            self.alert_text.setText("Укажите тех.карту для удаления")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)

    def cancel_added_document(self):
        self.enable_widget(True)
        self.add_document.setStyleSheet("QPushButton#add_document {\n"
                                        "    background-color: black; \n"
                                        "    color: white;\n"
                                        "    border-radius: 10px;\n"
                                        "}\n"
                                        "\n"
                                        "QPushButton#add_document:hover {\n"
                                        "    background-color: green;\n"
                                        "}")
        self.delete_document.setStyleSheet("QPushButton#delete_document {\n"
                                           "    background-color: black; \n"
                                           "    color: white;\n"
                                           "    border-radius: 10px;\n"
                                           "}\n"
                                           "\n"
                                           "QPushButton#delete_document:hover {\n"
                                           "    background-color: red;\n"
                                           "}")

    def add_file(self):
        self.window_add_document_name_document.clear()
        self.enable_widget(False)
        self.add_document.setStyleSheet("QPushButton#add_document {\n"
                                        "    background-color: rgb(115, 115, 115); \n"
                                        "    color: white;\n"
                                        "    border-radius: 10px;\n"
                                        "}\n"
                                        "\n"
                                        "QPushButton#add_document:hover {\n"
                                        "    background-color: green;\n"
                                        "}")
        self.delete_document.setStyleSheet("QPushButton#delete_document {\n"
                                           "    background-color: rgb(115, 115, 115); \n"
                                           "    color: white;\n"
                                           "    border-radius: 10px;\n"
                                           "}\n"
                                           "\n"
                                           "QPushButton#delete_document:hover {\n"
                                           "    background-color: red;\n"
                                           "}")

    def load_table(self):
        # Создаем список для хранения данных о файлах
        try:
            table_lst = []

            # Путь к каталогу с файлами
            directory = self.dir

            # Получаем список файлов в каталоге
            file_list = glob.glob(directory + "*.docx")

            # Выводим список файлов
            for file_path in file_list:
                # Получаем информацию о файле
                if not re.search(r"~\$", file_path):
                    file_name = os.path.basename(file_path)
                    creation_time = datetime.datetime.fromtimestamp(os.path.getctime(file_path)).strftime(
                        '%d.%m.%Y %H:%M:%S')
                    modification_time = datetime.datetime.fromtimestamp(os.path.getmtime(file_path)).strftime(
                        '%d.%m.%Y %H:%M:%S')
                    doc = Document(file_path)
                    core_props = doc.core_properties
                    authors = core_props.author
                    last_modified_by = core_props.last_modified_by
                    info_text = [file_name.split('.')[0], creation_time, modification_time, authors, last_modified_by]
                    table_lst.append(info_text)
            # Заполняем таблицу данными из списка
            self.table_docx.setRowCount(len(table_lst))
            for row_index, row_data in enumerate(table_lst):
                for col_index, col_data in enumerate(row_data):
                    item = QtWidgets.QTableWidgetItem(col_data)
                    self.table_docx.setItem(row_index, col_index, item)
        except Exception as ex:
            print(ex)

    def turn_document(self, row, col):
        self.row = self.table_docx.item(row, 0).text()

    def open_document(self, row, col):
        try:
            if col == 0:  # Проверяем, что клик произошел в столбце "Название"
                file_name_item = self.table_docx.item(row, 0)
                if file_name_item is not None:
                    file_name = file_name_item.text()
                    file_path = os.path.join(self.dir, file_name + ".docx")
                    self.doc_page = DocPage(self.main, file_path)
                    self.doc_page.setObjectName("doc_page")
                    self.main.stackedWidget.addWidget(self.doc_page)
                    self.main.stackedWidget.setCurrentWidget(self.doc_page)
                self.search_widget.clear()
        except Exception as ex:
            print(ex)

    def open_sidebar(self):
        self.home_content.setGeometry(QtCore.QRect(10, 10, 881, 861))
        self.table_docx.setGeometry(QtCore.QRect(20, 50, 851, 751))
        self.delete_document.setGeometry(QtCore.QRect(610, 820, 251, 31))
        self.alert.setGeometry(QtCore.QRect(40, 270, 821, 201))
        self.alert_title.setGeometry(QtCore.QRect(40, 270, 821, 61))
        self.alert_text.setGeometry(QtCore.QRect(40, 330, 821, 141))

    def close_sidebar(self):
        self.table_docx.setGeometry(QtCore.QRect(20, 50, 1001, 751))
        self.home_content.setGeometry(QtCore.QRect(10, 10, 1021, 861))
        self.delete_document.setGeometry(QtCore.QRect(750, 820, 251, 31))
        self.alert.setGeometry(QtCore.QRect(40, 270, 821, 201))
        self.alert_title.setGeometry(QtCore.QRect(40, 270, 961, 61))
        self.alert_text.setGeometry(QtCore.QRect(40, 330, 961, 141))
