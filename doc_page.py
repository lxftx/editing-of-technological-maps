import datetime

from PyQt5.QtCore import QTimer
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QWidget, QListWidget, QAbstractItemView

from alert import AlertMessage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import re
import os


class DocPage(QWidget, AlertMessage):
    def __init__(self, main, doc):
        super().__init__()
        self.doc_page = self
        self.main = main
        self.doc = doc
        self.timer = QTimer()
        self.row = 0 # переменная которая будет хранить строку при нажатии на таблицу
        self.content_info = [] # лист в котором будут собираться информация о тех.карте

        self.setWindowTitle("TechMap")
        self.setWindowIcon(QtGui.QIcon(self.get_path('images/icon', 'logo.png')))
        self.doc_content = QtWidgets.QLabel(self.doc_page)
        self.doc_content.setGeometry(QtCore.QRect(10, 10, 1021, 861))
        font = QtGui.QFont()
        font.setPointSize(30)
        font.setBold(True)
        font.setWeight(75)
        self.doc_content.setFont(font)
        self.doc_content.setStyleSheet("background-color:rgb(240, 240, 240);\n"
                                       "border-radius: 30px;\n"
                                       "color: rgb(255, 255, 17);")
        self.doc_content.setText("")
        self.doc_content.setObjectName("doc_content")

        self.open_list_table = QtWidgets.QPushButton(self.doc_page)
        self.open_list_table.setGeometry(QtCore.QRect(740, 20, 41, 31))

        icon = QtGui.QIcon()
        icon.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "open_table_in.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.open_list_table.setIcon(icon)
        self.open_list_table.setObjectName("open_list_table")
        self.open_list_table.clicked.connect(self.show_list_table)

        self.open_scroll = QtWidgets.QPushButton(self.doc_page)
        self.open_scroll.setGeometry(QtCore.QRect(850, 60, 171, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(10)
        self.open_scroll.setFont(font)


        icon1 = QtGui.QIcon()
        icon1.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "strelka_left.svg")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.open_scroll.setIcon(icon1)
        self.open_scroll.setObjectName("open_scroll")
        self.open_scroll.setCheckable(True)
        self.open_scroll.clicked.connect(self.unwrap_table)

        self.reset_button = QtWidgets.QPushButton(self.doc_page)
        self.reset_button.setGeometry(QtCore.QRect(794, 20, 41, 31))

        icon2 = QtGui.QIcon()
        icon2.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "reset_in.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.reset_button.setIcon(icon2)
        self.reset_button.setObjectName("reset_button")

        self.cursor = QtWidgets.QLabel(self.doc_page)
        self.cursor.setGeometry(QtCore.QRect(600, 60, 251, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.cursor.setFont(font)
        self.cursor.setStyleSheet("background-color:rgb(240, 240, 240);\n"
                                   "color: rgb(197, 197, 197);")
        self.cursor.setObjectName("cursor")
        self.cursor.setText(f"Курсор установлен на {self.row} строке")

        self.name_operation = QtWidgets.QLabel(self.doc_page)
        self.name_operation.setGeometry(QtCore.QRect(22, 70, 176, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.name_operation.setFont(font)
        self.name_operation.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.name_operation.setObjectName("name_operation")

        self.name_operation_box = QtWidgets.QComboBox(self.doc_page)
        self.name_operation_box.setGeometry(QtCore.QRect(200, 70, 231, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.name_operation_box.setFont(font)
        self.name_operation_box.setObjectName("name_operation_box")
        self.name_operation_box.addItems(['', 'Волочение', 'Отжиг'])

        self.oborudovanie = QtWidgets.QLabel(self.doc_page)
        self.oborudovanie.setGeometry(QtCore.QRect(33, 111, 108, 71))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.oborudovanie.setFont(font)
        self.oborudovanie.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.oborudovanie.setObjectName("oborudovanie")

        self.oborudovanie_edit = QtWidgets.QTextEdit(self.doc_page)
        self.oborudovanie_edit.setGeometry(QtCore.QRect(140, 110, 701, 71))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.oborudovanie_edit.setFont(font)
        self.oborudovanie_edit.setObjectName("oborudovanie_edit")

        self.label_diametr = QtWidgets.QLabel(self.doc_page)
        self.label_diametr.setGeometry(QtCore.QRect(30, 190, 811, 121))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.label_diametr.setFont(font)
        self.label_diametr.setStyleSheet("background-color:rgb(240, 240, 240); border: 1px solid black;")
        self.label_diametr.setText("")
        self.label_diametr.setObjectName("label_diametr")

        self.diametr = QtWidgets.QLabel(self.doc_page)
        self.diametr.setGeometry(QtCore.QRect(52, 210, 91, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.diametr.setFont(font)
        self.diametr.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.diametr.setObjectName("diametr")

        self.diametr_edit = QtWidgets.QLineEdit(self.doc_page)
        self.diametr_edit.setGeometry(QtCore.QRect(152, 210, 241, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.diametr_edit.setFont(font)
        self.diametr_edit.setObjectName("diametr_edit")
        self.diametr_edit.textChanged.connect(self.combining_diametr)

        self.diametr_do = QtWidgets.QLabel(self.doc_page)
        self.diametr_do.setGeometry(QtCore.QRect(52, 250, 131, 51))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.diametr_do.setFont(font)
        self.diametr_do.setStyleSheet("background-color:rgb(240, 240, 240); ")
        self.diametr_do.setObjectName("diametr_do")

        self.diametr_do_edit = QtWidgets.QLineEdit(self.doc_page)
        self.diametr_do_edit.setGeometry(QtCore.QRect(181, 261, 221, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.diametr_do_edit.setFont(font)
        self.diametr_do_edit.setObjectName("diametr_do_edit")
        self.diametr_do_edit.textChanged.connect(self.lock_combin)

        self.diametr_posle = QtWidgets.QLabel(self.doc_page)
        self.diametr_posle.setGeometry(QtCore.QRect(420, 250, 151, 51))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.diametr_posle.setFont(font)
        self.diametr_posle.setStyleSheet("background-color:rgb(240, 240, 240); ")
        self.diametr_posle.setObjectName("diametr_posle")

        self.diametr_posle_edit = QtWidgets.QLineEdit(self.doc_page)
        self.diametr_posle_edit.setGeometry(QtCore.QRect(580, 260, 231, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.diametr_posle_edit.setFont(font)
        self.diametr_posle_edit.setObjectName("diametr_posle_edit")
        self.diametr_do_edit.textChanged.connect(self.lock_combin)

        self.number_marshrut = QtWidgets.QLabel(self.doc_page)
        self.number_marshrut.setGeometry(QtCore.QRect(50, 330, 91, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.number_marshrut.setFont(font)
        self.number_marshrut.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.number_marshrut.setObjectName("number_marshrut")

        self.number_marshrut_edit = QtWidgets.QLineEdit(self.doc_page)
        self.number_marshrut_edit.setGeometry(QtCore.QRect(145, 330, 221, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.number_marshrut_edit.setFont(font)
        self.number_marshrut_edit.setObjectName("number_marshrut_edit")
        self.number_marshrut_edit.textChanged.connect(self.lock_heart_combine)

        self.steps_deformac = QtWidgets.QLabel(self.doc_page)
        self.steps_deformac.setGeometry(QtCore.QRect(420, 330, 172, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.steps_deformac.setFont(font)
        self.steps_deformac.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.steps_deformac.setObjectName("steps_deformac")

        self.steps_deformac_edit = QtWidgets.QLineEdit(self.doc_page)
        self.steps_deformac_edit.setGeometry(QtCore.QRect(600, 330, 231, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.steps_deformac_edit.setFont(font)
        self.steps_deformac_edit.setObjectName("steps_deformac_edit")
        self.steps_deformac_edit.textChanged.connect(self.lock_heart_combine)

        self.label_heattreatment = QtWidgets.QLabel(self.doc_page)
        self.label_heattreatment.setGeometry(QtCore.QRect(30, 380, 811, 201))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.label_heattreatment.setFont(font)
        self.label_heattreatment.setStyleSheet("background-color:rgb(240, 240, 240); border: 1px solid black;")
        self.label_heattreatment.setText("")
        self.label_heattreatment.setObjectName("label_heattreatment")

        self.heat_treatment_mode = QtWidgets.QLabel(self.doc_page)
        self.heat_treatment_mode.setGeometry(QtCore.QRect(50, 390, 171, 81))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.heat_treatment_mode.setFont(font)
        self.heat_treatment_mode.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.heat_treatment_mode.setObjectName("heat_treatment_mode")

        self.heat_treatment_mode_edit = QtWidgets.QTextEdit(self.doc_page)
        self.heat_treatment_mode_edit.setGeometry(QtCore.QRect(220, 390, 611, 81))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.heat_treatment_mode_edit.setFont(font)
        self.heat_treatment_mode_edit.setObjectName("heat_treatment_mode_edit")
        self.heat_treatment_mode_edit.textChanged.connect(self.combining_heat_treatment)

        self.heating_medium = QtWidgets.QLabel(self.doc_page)
        self.heating_medium.setGeometry(QtCore.QRect(50, 480, 101, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.heating_medium.setFont(font)
        self.heating_medium.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.heating_medium.setObjectName("heating_medium")

        self.heating_medium_edit = QtWidgets.QLineEdit(self.doc_page)
        self.heating_medium_edit.setGeometry(QtCore.QRect(160, 480, 211, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.heating_medium_edit.setFont(font)
        self.heating_medium_edit.setObjectName("heating_medium_edit")
        self.heating_medium_edit.textChanged.connect(self.lock_heart_combine)

        self.annealing_temperature = QtWidgets.QLabel(self.doc_page)
        self.annealing_temperature.setGeometry(QtCore.QRect(410, 480, 171, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.annealing_temperature.setFont(font)
        self.annealing_temperature.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.annealing_temperature.setObjectName("annealing_temperature")

        self.annealing_temperature_edit = QtWidgets.QLineEdit(self.doc_page)
        self.annealing_temperature_edit.setGeometry(QtCore.QRect(590, 480, 241, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.annealing_temperature_edit.setFont(font)
        self.annealing_temperature_edit.setObjectName("annealing_temperature_edit")
        self.annealing_temperature_edit.textChanged.connect(self.lock_heart_combine)

        self.time = QtWidgets.QLabel(self.doc_page)
        self.time.setGeometry(QtCore.QRect(50, 520, 81, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.time.setFont(font)
        self.time.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.time.setObjectName("time")

        self.time_edit = QtWidgets.QLineEdit(self.doc_page)
        self.time_edit.setGeometry(QtCore.QRect(140, 520, 231, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.time_edit.setFont(font)
        self.time_edit.setObjectName("time_edit")
        self.time_edit.textChanged.connect(self.lock_heart_combine)

        self.cooling_conditions = QtWidgets.QLabel(self.doc_page)
        self.cooling_conditions.setGeometry(QtCore.QRect(420, 520, 151, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.cooling_conditions.setFont(font)
        self.cooling_conditions.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.cooling_conditions.setObjectName("cooling_conditions")

        self.cooling_conditions_edit = QtWidgets.QLineEdit(self.doc_page)
        self.cooling_conditions_edit.setGeometry(QtCore.QRect(590, 520, 241, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.cooling_conditions_edit.setFont(font)
        self.cooling_conditions_edit.setObjectName("cooling_conditions_edit")
        self.cooling_conditions_edit.textChanged.connect(self.lock_heart_combine)

        self.note = QtWidgets.QLabel(self.doc_page)
        self.note.setGeometry(QtCore.QRect(50, 590, 91, 81))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.note.setFont(font)
        self.note.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.note.setObjectName("note")

        self.note_edit = QtWidgets.QTextEdit(self.doc_page)
        self.note_edit.setGeometry(QtCore.QRect(150, 590, 691, 81))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.note_edit.setFont(font)
        self.note_edit.setObjectName("note_edit")

        self.labe_column = QtWidgets.QLabel(self.doc_page)
        self.labe_column.setGeometry(QtCore.QRect(30, 680, 811, 121))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.labe_column.setFont(font)
        self.labe_column.setStyleSheet("background-color:rgb(240, 240, 240); border: 1px solid black;")
        self.labe_column.setText("")
        self.labe_column.setObjectName("labe_column")

        self.label = QtWidgets.QLabel(self.doc_page)
        self.label.setGeometry(QtCore.QRect(50, 690, 61, 61))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.label.setFont(font)
        self.label.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.label.setObjectName("label")

        self.lineEdit = QtWidgets.QLineEdit(self.doc_page)
        self.lineEdit.setGeometry(QtCore.QRect(120, 690, 711, 61))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.lineEdit.setFont(font)
        self.lineEdit.setObjectName("lineEdit")

        self.button_column = QtWidgets.QPushButton(self.doc_page)
        self.button_column.setGeometry(QtCore.QRect(50, 760, 781, 31))
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(10)
        self.button_column.setFont(font)
        icon3 = QtGui.QIcon()
        icon3.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "add_meger_column.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.button_column.setIcon(icon3)
        self.button_column.setObjectName("button_column")

        self.button_save_column = QtWidgets.QPushButton(self.doc_page)
        self.button_save_column.setGeometry(QtCore.QRect(20, 820, 181, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(10)
        self.button_save_column.setFont(font)
        icon4 = QtGui.QIcon()
        icon4.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "save_column.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.button_save_column.setIcon(icon4)
        self.button_save_column.setObjectName("button_save_column")

        self.button_add_column = QtWidgets.QPushButton(self.doc_page)
        self.button_add_column.setGeometry(QtCore.QRect(220, 820, 181, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(10)
        self.button_add_column.setFont(font)
        icon5 = QtGui.QIcon()
        icon5.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "add_column.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.button_add_column.setIcon(icon5)
        self.button_add_column.setObjectName("button_add_column")

        self.button_del_column = QtWidgets.QPushButton(self.doc_page)
        self.button_del_column.setGeometry(QtCore.QRect(420, 820, 181, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        self.button_del_column.setFont(font)
        icon6 = QtGui.QIcon()
        icon6.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "delete.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.button_del_column.setIcon(icon6)
        self.button_del_column.setObjectName("button_del_column")

        self.save_file = QtWidgets.QPushButton(self.doc_page)
        self.save_file.setGeometry(QtCore.QRect(20, 20, 181, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(10)
        self.save_file.setFont(font)
        icon7 = QtGui.QIcon()
        icon7.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "save_file.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.save_file.setIcon(icon7)
        self.save_file.setObjectName("save_file")

        self.open_file = QtWidgets.QPushButton(self.doc_page)
        self.open_file.setGeometry(QtCore.QRect(220, 20, 181, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(10)
        self.open_file.setFont(font)
        icon8 = QtGui.QIcon()
        icon8.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "open_file.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.open_file.setIcon(icon8)
        self.open_file.setObjectName("open_file")

        self.delete_file = QtWidgets.QPushButton(self.doc_page)
        self.delete_file.setGeometry(QtCore.QRect(420, 20, 181, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(10)
        self.delete_file.setFont(font)
        self.delete_file.setIcon(icon6)
        self.delete_file.setObjectName("delete_file")

        self.window = QtWidgets.QLabel(self.doc_page)
        self.window.setGeometry(QtCore.QRect(230, 110, 521, 601))
        self.window.setStyleSheet("border:1px solid black;")
        self.window.setText("")
        self.window.setObjectName("window")

        self.table_widget = QListWidget(self.doc_page)
        self.table_widget.setStyleSheet("font-size: 32px;")
        self.table_widget.itemDoubleClicked.connect(self.open_table)
        if self.main.user.posts == 'Технолог':
           self.table_widget.itemClicked.connect(self.enable_table)

        self.tables = QtWidgets.QScrollArea(self.doc_page)
        self.tables.setGeometry(QtCore.QRect(230, 190, 521, 471))
        self.tables.setWidgetResizable(True)
        self.tables.setObjectName("tables")
        self.scrollAreaWidgetContents1 = QtWidgets.QWidget()
        self.scrollAreaWidgetContents1.setGeometry(QtCore.QRect(0, 0, 521, 471))
        self.scrollAreaWidgetContents1.setObjectName("scrollAreaWidgetContents")
        self.tables.setWidget(self.scrollAreaWidgetContents1)
        self.tables.setWidget(self.table_widget)

        self.table_window = QtWidgets.QLabel(self.doc_page)
        self.table_window.setGeometry(QtCore.QRect(232, 115, 511, 71))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(14)
        self.table_window.setFont(font)
        self.table_window.setAlignment(QtCore.Qt.AlignCenter)
        self.table_window.setObjectName("table_window")

        self.button_add_table = QtWidgets.QPushButton(self.doc_page)
        self.button_add_table.setGeometry(QtCore.QRect(240, 670, 241, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.button_add_table.setFont(font)
        self.button_add_table.setStyleSheet("QPushButton#button_add_table {\n"
                                            "    background-color: black; \n"
                                            "    color: white;\n"
                                            "    border-radius: 10px;\n"
                                            "}\n"
                                            "\n"
                                            "QPushButton#button_add_table:hover {\n"
                                            "    background-color: green;\n"
                                            "}")
        self.button_add_table.setObjectName("button_add_table")
        self.button_add_table.clicked.connect(self.add_new_table)

        self.button_del_table = QtWidgets.QPushButton(self.doc_page)
        self.button_del_table.setGeometry(QtCore.QRect(490, 670, 251, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.button_del_table.setFont(font)
        self.button_del_table.setObjectName("button_del_table")
        self.button_del_table.clicked.connect(self.delele_table)

        self.info_window = QtWidgets.QLabel(self.doc_page)
        self.info_window.setGeometry(QtCore.QRect(280, 150, 401, 459))
        self.info_window.setStyleSheet("background-color:rgb(240, 240, 240);\n"
                                       "border: 1px solid black;\n"
                                       "border-radius: 15px;\n"
                                       "color: rgb(255, 255, 17);")
        self.info_window.setText("")
        self.info_window.setObjectName("info_window")
        self.info_window.setVisible(False)

        self.director = QtWidgets.QLabel(self.doc_page)
        self.director.setGeometry(QtCore.QRect(300, 168, 361, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(16)
        self.director.setFont(font)
        self.director.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.director.setAlignment(QtCore.Qt.AlignCenter)
        self.director.setObjectName("director")
        self.director.setVisible(False)

        self.director_edit = QtWidgets.QLineEdit(self.doc_page)
        self.director_edit.setGeometry(QtCore.QRect(290, 198, 381, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.director_edit.setFont(font)
        self.director_edit.setText("")
        self.director_edit.setObjectName("director_edit")
        self.director_edit.setVisible(False)

        self.number_card = QtWidgets.QLabel(self.doc_page)
        self.number_card.setGeometry(QtCore.QRect(300, 238, 361, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(16)
        self.number_card.setFont(font)
        self.number_card.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.number_card.setAlignment(QtCore.Qt.AlignCenter)
        self.number_card.setObjectName("number_card")
        self.number_card.setVisible(False)

        self.number_card_edit = QtWidgets.QLineEdit(self.doc_page)
        self.number_card_edit.setGeometry(QtCore.QRect(290, 268, 381, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.number_card_edit.setFont(font)
        self.number_card_edit.setObjectName("number_card_edit")
        self.number_card_edit.setVisible(False)

        self.splav = QtWidgets.QLabel(self.doc_page)
        self.splav.setGeometry(QtCore.QRect(300, 308, 361, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(16)
        self.splav.setFont(font)
        self.splav.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.splav.setAlignment(QtCore.Qt.AlignCenter)
        self.splav.setObjectName("volochen_item")
        self.splav.setVisible(False)

        self.splav_edit = QtWidgets.QLineEdit(self.doc_page)
        self.splav_edit.setGeometry(QtCore.QRect(290, 338, 381, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.splav_edit.setFont(font)
        self.splav_edit.setText("")
        self.splav_edit.setObjectName("volochen_item_edit")
        self.splav_edit.setVisible(False)

        self.gost = QtWidgets.QLabel(self.doc_page)
        self.gost.setGeometry(QtCore.QRect(300, 378, 361, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(16)
        self.gost.setFont(font)
        self.gost.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.gost.setAlignment(QtCore.Qt.AlignCenter)
        self.gost.setObjectName("gost")
        self.gost.setVisible(False)

        self.gost_edit = QtWidgets.QLineEdit(self.doc_page)
        self.gost_edit.setGeometry(QtCore.QRect(290, 408, 381, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.gost_edit.setFont(font)
        self.gost_edit.setObjectName("gost_edit")
        self.gost_edit.setVisible(False)

        self.sto = QtWidgets.QLabel(self.doc_page)
        self.sto.setGeometry(QtCore.QRect(300, 448, 361, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(16)
        self.sto.setFont(font)
        self.sto.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.sto.setAlignment(QtCore.Qt.AlignCenter)
        self.sto.setObjectName("sto")
        self.sto.setVisible(False)

        self.sto_edit = QtWidgets.QLineEdit(self.doc_page)
        self.sto_edit.setGeometry(QtCore.QRect(290, 478, 381, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.sto_edit.setFont(font)
        self.sto_edit.setObjectName("sto_edit")
        self.sto_edit.setVisible(False)

        self.vzamen = QtWidgets.QLabel(self.doc_page)
        self.vzamen.setGeometry(QtCore.QRect(300, 518, 361, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(16)
        self.vzamen.setFont(font)
        self.vzamen.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.vzamen.setAlignment(QtCore.Qt.AlignCenter)
        self.vzamen.setObjectName("vzamen")
        self.vzamen.setVisible(False)

        self.vzamen_edit = QtWidgets.QLineEdit(self.doc_page)
        self.vzamen_edit.setGeometry(QtCore.QRect(290, 548, 381, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.vzamen_edit.setFont(font)
        self.vzamen_edit.setObjectName("vzamen_edit")
        self.vzamen_edit.setVisible(False)

        self.button_info_window = QtWidgets.QPushButton(self.doc_page)
        self.button_info_window.setGeometry(QtCore.QRect(850, 20, 171, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(10)
        self.button_info_window.setFont(font)
        icon9 = QtGui.QIcon()
        icon9.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "info_window.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.button_info_window.setIcon(icon9)
        self.button_info_window.setIconSize(QtCore.QSize(25, 25))
        self.button_info_window.setObjectName("button_info_window")
        self.button_info_window.clicked.connect(lambda: self.open_info_window(True))

        self.alert = QtWidgets.QLabel(self.doc_page)
        self.alert.setGeometry(QtCore.QRect(40, 270, 821, 201))
        self.alert.setStyleSheet("border-radius: 30px;")
        self.alert.setObjectName("alert_info")
        self.alert.setVisible(False)

        self.alert_title = QtWidgets.QLabel(self.doc_page)
        self.alert_title.setGeometry(QtCore.QRect(40, 270, 961, 61))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(20)
        font.setBold(True)
        font.setWeight(75)
        self.alert_title.setFont(font)
        self.alert_title.setStyleSheet("border-top-left-radius: 30px;\n"
                                       "border-top-right-radius: 30px;\n"
                                       "background-color: rgb(227, 227, 227); \n"
                                       "color: rgb(179, 179, 179);")
        self.alert_title.setAlignment(QtCore.Qt.AlignCenter)
        self.alert_title.setObjectName("alert_info_title")
        self.alert_title.setText("Предупреждение")
        self.alert_title.setVisible(False)

        self.alert_text = QtWidgets.QLabel(self.doc_page)
        self.alert_text.setGeometry(QtCore.QRect(40, 330, 961, 141))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(15)
        font.setUnderline(True)
        self.alert_text.setFont(font)
        self.alert_text.setStyleSheet("border-bottom-left-radius: 30px;\n"
                                      "border-bottom-right-radius: 30px;\n"
                                      "color: rgb(162, 162, 162);")
        self.alert_text.setAlignment(QtCore.Qt.AlignCenter)
        self.alert_text.setObjectName("alert_info_text")
        self.alert_text.setVisible(False)

        self.table = QtWidgets.QTableWidget(self.doc_page)
        self.table.setGeometry(QtCore.QRect(850, 110, 171, 701))
        self.table.setTabletTracking(False)
        self.table.setObjectName("table")
        self.table.setColumnCount(12)  # Увеличиваем количество столбцов на 1
        self.table.setRowCount(0)
        self.table.setHorizontalHeaderLabels(["№п/п", "Наименование операций", "Оборудование", "До волочения", "После волочения", '№ маршрута', 'Степень деформации, %', 'Среда нагрева', 'Температура отжига, ̊С', 'Время, мин', 'Условия охлаждения', 'Примечание'])
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        if self.main.user.posts == 'Технолог':
            self.table.cellClicked.connect(self.on_row)
        self.table.cellDoubleClicked.connect(self.turn_item_in_table)


        self.director.setText("Директор:")
        self.gost.setText("по ГОСТ:")
        self.vzamen.setText("Взамен")
        self.button_add_column.setText("Добавить колонку")
        self.number_card.setText("№ технологической карты:")
        self.splav.setText("сплав:")
        self.vzamen.setText('Взамен:')
        self.diametr.setText("Диаметр, мм:")
        self.note.setText("Примечание:")
        self.oborudovanie.setText("Оборудование:")
        self.diametr_do.setText("Диаметр \nдо волочения, мм:")
        self.sto.setText("и СТО:")
        self.number_marshrut.setText("№ маршрута:")
        self.name_operation.setText("Наименование операций:")
        self.open_file.setText("Открыть файл")
        self.open_scroll.setText("Открыть")
        self.steps_deformac.setText("Степень деформации, %:")
        self.diametr_posle.setText("Диаметр \nпосле волочения, мм:")
        self.delete_file.setText("Удалить файл")
        self.heat_treatment_mode.setText("Режим термообработки:")
        self.heating_medium.setText("Среда нагрева:")
        self.annealing_temperature.setText("Температура отжига, *С:")
        self.cooling_conditions.setText("Условия охлаждения:")
        self.time.setText("Время, мин:")
        self.label.setText("Для чего:")
        self.button_column.setText("Добавить")
        self.button_save_column.setText("Сохранить колонку")
        self.button_del_column.setText("Удалить колонку")
        self.save_file.setText("Сохранить файл")
        self.table_window.setText("Выберите таблицу для редактирования")
        self.button_add_table.setText("Добавить таблицу")
        self.button_del_table.setText("Удалить таблицу")
        self.button_info_window.setText("Инфо. тех.карта")

        self.button_save_column.clicked.connect(self.save_column)
        self.button_del_column.clicked.connect(self.del_column)
        self.button_column.clicked.connect(self.add_merger_column)
        self.button_add_column.clicked.connect(self.add_column_item)
        self.save_file.clicked.connect(self.save_file_item)
        self.open_file.clicked.connect(self.open_file_docx)
        self.delete_file.clicked.connect(self.del_file_docx)
        self.reset_button.clicked.connect(self.clear_widget)
        self.load_table_item()

        self.reset_button.installEventFilter(self)
        self.open_list_table.installEventFilter(self)

    def mousePressEvent(self, event):
        self.open_info_window(False)
        super().mousePressEvent(event)

    def closeEvent(self, event):
        event.accept()

    def eventFilter(self, watched, event):
        icon = QtGui.QIcon()
        if (event.type() == QtCore.QEvent.Enter) and (watched == self.reset_button):
            icon.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "reset_in.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
            self.reset_button.setIcon(icon)
        elif (event.type() == QtCore.QEvent.Leave) and (watched == self.reset_button):
            icon.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "reset_out.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
            self.reset_button.setIcon(icon)
        if (event.type() == QtCore.QEvent.Enter) and (watched == self.open_list_table):
            icon.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "open_table_in.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
            self.open_list_table.setIcon(icon)
        elif (event.type() == QtCore.QEvent.Leave) and (watched == self.open_list_table):
            icon.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "open_table_out.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
            self.open_list_table.setIcon(icon)
        return super().eventFilter(watched, event)

    def get_path(self, dir, name):
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), dir, name)

    # Функция, которая загружает в список сущетсвующих таблиц
    def load_table_item(self):
        self.table_widget.clear()
        doc = Document(self.doc)
        for i, table in enumerate(doc.tables):
            if all(len(row.cells) == 12 for row in table.rows):
                item = QtWidgets.QListWidgetItem(f'Таблица - {str(i + 1)}')
                item.setTextAlignment(QtCore.Qt.AlignCenter)
                self.table_widget.addItem(item)
        # Установка не рабочих виджетов
        self.widget_enable_true_or_false(False)
        self.set_stylesheet_widget("QPushButton {\n"
                                       "    background-color: rgb(115, 115, 115); \n"
                                       "    color: black;\n"
                                       "    border-radius: 10px;\n"
                                       "}\n"
                                       "\n"
                                       "QPushButton:hover {\n"
                                       "    background-color: green;\n"
                                       "}", "QPushButton {\n"
                                       "    background-color: rgb(115, 115, 115); \n"
                                       "    color: black;\n"
                                       "    border-radius: 10px;\n"
                                       "}\n"
                                       "\n"
                                       "QPushButton:hover {\n"
                                       "    background-color: red;\n"
                                       "}", False)

    # Выбор таблицы
    def enable_table(self, item):
        self.table_id = int(item.text()[-1]) - 1 # берем число из текста, делаем int и вычитаем 1 так как прибавляется 1
        self.row_table = 0 if self.table_id else 2  # Чтение будет производится после какой строки, если таблица 0, то с 2, если нет, то с 0
        if not self.table_widget.count() == 1:
            self.button_del_table.setEnabled(True)
            self.button_del_table.setStyleSheet("QPushButton#button_del_table {\n"
                                                "    background-color: black; \n"
                                                "    color: white;\n"
                                                "    border-radius: 10px;\n"
                                                "}\n"
                                                "\n"
                                                "QPushButton#button_del_table:hover {\n"
                                                "    background-color: red;\n"
                                                "}")
        else:
            self.button_del_table.setEnabled(False)
            self.button_del_table.setStyleSheet("QPushButton#button_del_table {\n"
                                                "    background-color: rgb(115, 115, 115);; \n"
                                                "    color: white;\n"
                                                "    border-radius: 10px;\n"
                                                "}\n"
                                                "\n"
                                                "QPushButton#button_del_table:hover {\n"
                                                "    background-color: red;\n"
                                                "}")

    # Создание новой таблицы
    def add_new_table(self):
        try:
            doc = Document(self.doc)
            doc.add_section(WD_SECTION_START.NEW_PAGE)
            table = doc.add_table(rows=1, cols=12)

            # Применяем стиль "Table Grid" к таблице
            table.style = "Table Grid"

            for i in range(12):
                cell = table.cell(0, i)
                cell.text = str(i + 1)
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    run = paragraph.runs[0]
                    run.bold = True

            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(10)  # Используем Pt для установки размера шрифта
                        # Выравниваем текст по центру ячеек
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.save(self.doc)
            self.table_widget.clear()
            self.load_table_item()
            self.modific_doc_file()
        except PermissionError:
            self.show_alert()
            self.alert_text.setText(f"Файл был открыт!\nЗайкроте его и снова выполните создание таблицы!")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)


    # Функция, которая удаляет таблицу
    def delele_table(self):
        try:
            if not self.table_widget.count() == 1:
                # Удаляем таблицу с файла
                doc = Document(self.doc)
                table_to_delete = doc.tables[self.table_id]
                table_to_delete._element.getparent().remove(table_to_delete._element)
                doc.save(self.doc)

                # Обновляем список таблиц и имя пользователя который изменил файл
                self.table_widget.clear()
                self.load_table_item()
                self.modific_doc_file()

                # При удалении курсор будет смещаться на предыдущую таблицу.
                self.table_id -= 1
            else:
                self.button_del_table.setEnabled(False)
                self.button_del_table.setStyleSheet("QPushButton#button_del_table {\n"
                                                    "    background-color: rgb(115, 115, 115);; \n"
                                                    "    color: white;\n"
                                                    "    border-radius: 10px;\n"
                                                    "}\n"
                                                    "\n"
                                                    "QPushButton#button_del_table:hover {\n"
                                                    "    background-color: red;\n"
                                                    "}")
        except PermissionError:
            self.show_alert()
            self.alert_text.setText("Файл был открыт!\nЗайкроте его и снова выполните удаление таблицы!")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)

    # Функция, которая открывает нынешний файл
    def open_file_docx(self):
        os.startfile(self.doc)

    # Функция, которая удаляет нынешний файл
    def del_file_docx(self):
        try:
            os.remove(self.doc)
            self.main.home_page.load_table()
            self.main.stackedWidget.setCurrentWidget(self.main.home_page)
        except PermissionError:
            self.show_alert()
            self.alert_text.setText(f"Файл был открыт!\nЗайкроте его и снова выполните удаление тех.карты!")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)

    # Функция, которая сохраняет все изменения в файле
    def save_file_item(self):
        try:
            self.save_info()
            old_filename = self.doc
            directory = os.path.dirname(old_filename)
            new_filename = f"{directory}/{self.number_card_edit.text()}.docx"
            self.doc = new_filename
            os.rename(old_filename, new_filename)
            self.modific_doc_file()
        except PermissionError:
            self.show_alert()
            self.alert_text.setText(f"Файл был открыт!\nЗайкроте его и снова выполните сохранение тех.карты!")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(3000)
        except Exception as ex:
            self.show_alert()
            self.alert_text.setText(f"Системная Ошибка! {ex}")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)

    # Функция, которая сохраняет информацию о тех. карте.
    def save_info(self):
        try:
            doc = Document(self.doc)
            # Замена текста "NUMBER_CARD" на другой текст и применение стилей
            for paragraph in doc.paragraphs:
                if self.content_info[0] in paragraph.text:
                    # Замена текста
                    paragraph.text = paragraph.text.replace(self.content_info[0], self.number_card_edit.text())
                    # Применение стилей
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Times New Roman'
                        font.size = Pt(11)
                        font.bold = True

            # Вставка текущего года вместо "YEAR" и применение стилей
            current_year = str(datetime.datetime.now().year)
            for paragraph in doc.paragraphs:
                if re.search(r'_([0-9]+) г\.', self.content_info[5]).group(1) in re.search(r'_([0-9]+) г\.',
                                                                                      paragraph.text).group(
                        1) if re.search(r'_([0-9]+) г\.', paragraph.text) else None:
                    # Замена текста
                    paragraph.text = paragraph.text.replace(re.search(r'_([0-9]+) г\.', self.content_info[5]).group(1),
                                                            current_year)
                elif 'YEAR' in paragraph.text:
                    # Получаем текущий год
                    current_year = datetime.datetime.now().year
                    paragraph.text = re.sub(r'YEAR', str(current_year), paragraph.text)
                elif re.search(r'_{2,}(.+)', self.content_info[4]).group(1) in re.search(r'_{2,}(.+)',
                                                                                    self.content_info[4]).group(
                        1) if re.search(r'_{2,}(.+)', paragraph.text) else None:
                    paragraph.text = paragraph.text.replace(re.search(r'_{2,}(.+)', self.content_info[4]).group(1),
                                                            self.director_edit.text())
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Times New Roman'
                    font.size = Pt(10)
                    font.bold = True

            # Маппинг ключевых слов и значений для вставки и применение стилей
            keyword_value_map = {
                re.search(r'сплава\s+(.*?)\s+по', self.content_info[8]).group(1): self.splav_edit.text(),
                re.search(rf'\b{re.escape("ГОСТ")}\s+(.*?)\s+{re.escape("и")}', self.content_info[8]).group(
                    1): self.gost_edit.text(),
                re.search(rf'\b{re.escape("СТО")}\s+(.*)$', self.content_info[8]).group(1): self.sto_edit.text(),
                re.search(rf'\b{re.escape("Взамен")}\s+(.*?)$', self.content_info[10]).group(1): self.vzamen_edit.text()
            }
            for keyword, value in keyword_value_map.items():
                for paragraph in doc.paragraphs:
                    if keyword in paragraph.text.split():
                        # Замена ключевого слова на значение
                        paragraph.text = paragraph.text.replace(keyword, value)
                        # Применение стилей для значения
                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True

            doc.save(self.doc)
            self.modific_doc_file()
        except PermissionError:
            self.show_alert()
            self.alert_text.setText(f"Файл был открыт!\nЗайкроте его и снова выполните сохранение файла!")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(3000)

    # Функция, которая изменяем имя изменившего файл
    def modific_doc_file(self):
        try:
            doc = Document(self.doc)
            core_props = doc.core_properties
            core_props.last_modified_by = f'{self.main.user.surname} {self.main.user.name}'
            doc.save(self.doc)
        except Exception as e:
            self.show_alert()
            self.alert_text.setText(f"Системная Ошибка!\nПроизошла ошибка при изменении имени пользователя: {e}")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)

    # Функция которая открывает выбранную таблицу
    def open_table(self):
        self.clear_widget()
        self.widget_enable_true_or_false(True)
        self.hide_table_window(False)
        self.set_stylesheet_widget("QPushButton {\n"
                                            "    background-color: black; \n"
                                            "    color: white;\n"
                                            "    border-radius: 10px;\n"
                                            "}\n"
                                            "\n"
                                            "QPushButton:hover {\n"
                                            "    background-color: green;\n"
                                            "}", "QPushButton {\n"
                                            "    background-color: black; \n"
                                            "    color: white;\n"
                                            "    border-radius: 10px;\n"
                                            "}\n"
                                            "\n"
                                            "QPushButton:hover {\n"
                                            "    background-color: red;\n"
                                            "}", True)
        # Читаем данные из файла
        self.read_file_table()

    # Функция по чтению данных из таблиц
    def read_file_table(self):
        self.table.setRowCount(0)  # Сбрасывает кол-во строк в открывающей таблице
        doc = Document(self.doc)
        table = doc.tables[self.table_id]  # указывается какой номер таблицы
        table_items = []    # лист с уникальными данными
        table_content = []  # лист в котором буду содержаться при чтении из файла
        # Читаем содержимое таблицы
        for row_id, row in enumerate(table.rows):
            if row_id > self.row_table:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_content.append(row_data)
        # Перебираем и делаем уникальный вид
        print(table_content)
        for row in table_content:
            if len(set(row)) == 1 and '-' not in row:
                table_items.append([row[0]])
            else:
                unique_arr = []  # Создаем список для уникальных элементов в текущей строке
                seen_3_5 = set()
                seen_5_11 = set()
                for idx, item in enumerate(row):
                    if idx in [3, 4] and item in row[3:5]:
                        if item != '-' and item in seen_3_5:
                            continue
                        else:
                            unique_arr.append(item)
                            seen_3_5.add(item)
                    elif idx in [5, 6, 7, 8, 9, 10] and item in row[5:11]:
                        if item != '-' and item in seen_5_11:
                            continue
                        else:
                            unique_arr.append(item)
                            seen_5_11.add(item)
                    else:
                        unique_arr.append(item)
                table_items.append(unique_arr)
        self.table.setRowCount(len(table_items))    # Устанавливаем кол-во строк в таблице
        self.table.setColumnCount(12)               # Устанаваливаем кол-во колонок в таблице
        for row_index, row_data in enumerate(table_items):
            match len(row_data):
                case 1:
                    self.table.setSpan(row_index, 0, 1, 12)
                case 6:
                    self.table.setSpan(row_index, 3, 1, 2)
                    self.table.setSpan(row_index, 5, 1, 6)
                case 7:
                    self.table.setSpan(row_index, 5, 1, 6)
                case 11:
                    self.table.setSpan(row_index, 3, 1, 2)
                case _:
                    print('Ничего не объединяем')
            for row_index, row_data in enumerate(table_content):
                for col_index, col_data in enumerate(row_data):
                    item = QtWidgets.QTableWidgetItem(col_data)
                    item.setTextAlignment(0x0004 | 0x0080)
                    self.table.setItem(row_index, col_index, item)
        self.info()

    # Функция, которая считывает данные с тех.карты
    def info(self):
        doc = Document(self.doc)
        # Создаем пустой список для хранения текста и других элементов
        # Читаем содержимое каждого элемента в документе
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Обрабатываем параграфы
                self.content_info.append(element.text)
            elif element.tag.endswith('tbl'):  # Обрабатываем таблицы
                table_content = []
                for row in element.findall('.//tr'):
                    table_row = []
                    for cell in row.findall('.//tc'):
                        cell_text = ''
                        for paragraph in cell.findall('.//p'):
                            cell_text += paragraph.text + '\n'
                        table_row.append(cell_text.strip())
                    table_content.append(table_row)
                self.content_info.append(table_content)

        if len(self.content_info):
            self.number_card_edit.setText(self.content_info[0])
            match = re.search(r'_{2,}(.+)', self.content_info[4])
            if match:
                self.director_edit.setText(match.group(1))
            match = re.search(r'сплава\s+(.*?)\s+по', self.content_info[8])
            if match:
                self.splav_edit.setText(match.group(1))
            match = re.search(rf'\b{re.escape("ГОСТ")}\s+(.*?)\s+{re.escape("и")}', self.content_info[8])
            if match:
                self.gost_edit.setText(match.group(1))
            match = re.search(rf'\b{re.escape("СТО")}\s+(.*)$', self.content_info[8])
            if match:
                self.sto_edit.setText(match.group(1))
            match = re.search(rf'\b{re.escape("Взамен")}\s+(.*?)$', self.content_info[10])
            if match:
                self.vzamen_edit.setText(match.group(1))

    # Функция, которая присваивает элемент self.row
    def on_row(self, row):
        self.row = self.row_table + row + 1
        self.cursor.setText(f"Курсор установлен на {row + 1} строке")


    # Функция которая вызывыется, путем двумя нажатиями на элемент в таблице
    def turn_item_in_table(self, row):
        self.clear_widget() # Чистим все виджеты, после того, как хотим их перенести на виджеты
        row_data = []
        table_items = []
        for column in range(self.table.columnCount()):
            item = self.table.item(row, column)
            if item is not None:
                row_data.append(item.text())

        unique_arr = []
        if len(set(row_data)) == 1 and '-' not in row_data:
            unique_arr.append(row_data[0])
        else:
            seen_3_5 = set()
            seen_5_11 = set()
            for idx, item in enumerate(row_data):
                if idx in [3, 4] and item in row_data[3:5]:
                    if item != '-' and item in seen_3_5:
                        continue
                    else:
                        unique_arr.append(item)
                        seen_3_5.add(item)
                elif idx in [5, 6, 7, 8, 9, 10] and item in row_data[5:11]:
                    if item != '-' and item in seen_5_11:
                        continue
                    else:
                        unique_arr.append(item)
                        seen_5_11.add(item)
                else:
                    unique_arr.append(item)
        match len(unique_arr):
            case 1:
                self.lineEdit.setText(row_data[0])
            case 6:
                self.name_operation_box.setCurrentText(unique_arr[1]), self.oborudovanie_edit.setText(unique_arr[2])
                self.diametr_edit.setText(unique_arr[3]), self.heat_treatment_mode_edit.setText(unique_arr[4])
                self.note_edit.setText(unique_arr[5])
            case 7:
                self.name_operation_box.setCurrentText(unique_arr[1]), self.oborudovanie_edit.setText(unique_arr[2])
                self.diametr_do_edit.setText(unique_arr[3]), self.diametr_posle_edit.setText(unique_arr[4])
                self.heat_treatment_mode_edit.setText(unique_arr[5]), self.note_edit.setText(unique_arr[6])
            case 11:
                self.name_operation_box.setCurrentText(unique_arr[1]), self.oborudovanie_edit.setText(unique_arr[2])
                self.diametr_edit.setText(unique_arr[3]), self.number_marshrut_edit.setText(unique_arr[4])
                self.steps_deformac_edit.setText(unique_arr[5]), self.heating_medium_edit.setText(unique_arr[6])
                self.annealing_temperature_edit.setText(unique_arr[7]), self.time_edit.setText(unique_arr[8])
                self.cooling_conditions_edit.setText(unique_arr[9]), self.note_edit.setText(unique_arr[10])
            case 12:
                self.name_operation_box.setCurrentText(unique_arr[1]), self.oborudovanie_edit.setText(unique_arr[2])
                self.diametr_do_edit.setText(unique_arr[3]), self.diametr_posle_edit.setText(unique_arr[4])
                self.number_marshrut_edit.setText(unique_arr[5]) , self.steps_deformac_edit.setText(unique_arr[6])
                self.heating_medium_edit.setText(unique_arr[7]), self.annealing_temperature_edit.setText(unique_arr[8])
                self.time_edit.setText(unique_arr[9]), self.cooling_conditions_edit.setText(unique_arr[10])
                self.note_edit.setText(unique_arr[11])
            case _:
                print(f"Sorry, {len(unique_arr)}")
        self.row = row + self.row_table + 1

    def read_text_widgets_and_write_file(self, new_row):
        items = ['idx',
                 self.name_operation_box.currentText(),
                 self.oborudovanie_edit.toPlainText() if self.oborudovanie_edit.toPlainText().strip() else '-']
        list_widgets = [self.diametr_edit, self.diametr_do_edit, self.diametr_posle_edit, self.heat_treatment_mode_edit,
                        self.number_marshrut_edit, self.steps_deformac_edit, self.heating_medium_edit, self.annealing_temperature_edit
                        , self.time_edit, self.cooling_conditions_edit]
        for widget in list_widgets:
            if widget.isEnabled():
                try:
                    items.append(widget.text() if widget.text().strip() else '-')
                except Exception:
                    items.append(widget.toPlainText() if widget.toPlainText().strip() else '-')
            else:
                items.append("")
        items.append(self.note_edit.toPlainText() if self.note_edit.toPlainText().strip() else '-')
        print(items)
        items = len([x for x in items if len(x)])
        match items:
            case 12:
                for cell, text in zip(new_row.cells,
                                      ['idx',
                                       self.name_operation_box.currentText() if self.name_operation_box.currentText().strip() else self.name_operation_box.setCurrentIndex(
                                           0),
                                       self.oborudovanie_edit.toPlainText().strip() if self.oborudovanie_edit.toPlainText().strip() else '-',
                                       self.diametr_do_edit.text().strip() if self.diametr_do_edit.text().strip() else '-',
                                       self.diametr_posle_edit.text().strip() if self.diametr_posle_edit.text().strip() else '-',
                                       self.number_marshrut_edit.text().strip() if self.number_marshrut_edit.text().strip() else '-',
                                       self.steps_deformac_edit.text().strip() if self.steps_deformac_edit.text().strip() else '-',
                                       self.heating_medium_edit.text().strip() if self.heating_medium_edit.text().strip() else '-',
                                       self.annealing_temperature_edit.text().strip() if self.annealing_temperature_edit.text().strip() else '-',
                                       self.time_edit.text().strip() if self.time_edit.text().strip() else '-',
                                       self.cooling_conditions_edit.text().strip() if self.cooling_conditions_edit.text().strip() else '-',
                                       self.note_edit.toPlainText().strip() if self.note_edit.toPlainText().strip() else '-']):
                    cell.text = text
                    paragraph = cell.paragraphs[0]
                    # Выравниваем текст по центру
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Устанавливаем шрифт, стиль и размер шрифта для текста в объединенной ячейке
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
            case 11:
                new_row.cells[3].merge(new_row.cells[4])
                for cell, text in zip(new_row.cells,
                                      ['idx',
                                       self.name_operation_box.currentText() if self.name_operation_box.currentText().strip() else self.name_operation_box.setCurrentIndex(
                                           0),
                                       self.oborudovanie_edit.toPlainText().strip() if self.oborudovanie_edit.toPlainText().strip() else '-',
                                       '',
                                       self.diametr_edit.text().strip() if self.diametr_edit.text().strip() else '-',
                                       self.number_marshrut_edit.text().strip() if self.number_marshrut_edit.text().strip() else '-',
                                       self.steps_deformac_edit.text().strip() if self.steps_deformac_edit.text().strip() else '-',
                                       self.heating_medium_edit.text().strip() if self.heating_medium_edit.text().strip() else '-',
                                       self.annealing_temperature_edit.text().strip() if self.annealing_temperature_edit.text().strip() else '-',
                                       self.time_edit.text().strip() if self.time_edit.text().strip() else '-',
                                       self.cooling_conditions_edit.text().strip() if self.cooling_conditions_edit.text().strip() else '-',
                                       self.note_edit.toPlainText().strip() if self.note_edit.toPlainText().strip() else '-']):
                    cell.text = text
                    paragraph = cell.paragraphs[0]
                    # Выравниваем текст по центру
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Устанавливаем шрифт, стиль и размер шрифта для текста в объединенной ячейке
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
            case 6:
                new_row.cells[3].merge(new_row.cells[4])
                new_row.cells[5].merge(new_row.cells[10])
                for cell, text in zip(new_row.cells,
                                      ['idx',
                                       self.name_operation_box.currentText() if self.name_operation_box.currentText().strip() else self.name_operation_box.setCurrentIndex(
                                           0),
                                       self.oborudovanie_edit.toPlainText().strip() if self.oborudovanie_edit.toPlainText().strip() else '-',
                                       '',
                                       self.diametr_edit.text().strip() if self.diametr_edit.text().strip() else '-',
                                       '',
                                       '',
                                       '',
                                       '',
                                       '',
                                       self.heat_treatment_mode_edit.toPlainText().strip() if self.heat_treatment_mode_edit.toPlainText().strip() else '-',
                                       self.note_edit.toPlainText().strip() if self.note_edit.toPlainText().strip() else '-']):
                    cell.text = text
                    paragraph = cell.paragraphs[0]
                    # Выравниваем текст по центру
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Устанавливаем шрифт, стиль и размер шрифта для текста в объединенной ячейке
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
            case 7:
                new_row.cells[5].merge(new_row.cells[10])
                for cell, text in zip(new_row.cells,
                                      ['idx',
                                       self.name_operation_box.currentText() if self.name_operation_box.currentText().strip() else self.name_operation_box.setCurrentIndex(
                                           0),
                                       self.oborudovanie_edit.toPlainText().strip() if self.oborudovanie_edit.toPlainText().strip() else '-',
                                       self.diametr_do_edit.text().strip() if self.diametr_do_edit.text().strip() else '-',
                                       self.diametr_posle_edit.text().strip() if self.diametr_posle_edit.text().strip() else '-',
                                       '',
                                       '',
                                       '',
                                       '',
                                       '',
                                       self.heat_treatment_mode_edit.toPlainText().strip() if self.heat_treatment_mode_edit.toPlainText().strip() else '-',
                                       self.note_edit.toPlainText().strip() if self.note_edit.toPlainText().strip() else '-']):
                    cell.text = text
                    paragraph = cell.paragraphs[0]
                    # Выравниваем текст по центру
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Устанавливаем шрифт, стиль и размер шрифта для текста в объединенной ячейке
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

    # Функция, которая сохраняет текст в таблицу
    def save_column(self):
        try:
            if self.row:
                doc = Document(self.doc)
                row_to_remove_index = self.row
                table = doc.tables[self.table_id]
                if 0 <= row_to_remove_index < len(table.rows):
                    row_to_remove = table.rows[row_to_remove_index]
                    parent = row_to_remove._element.getparent()
                    parent.remove(row_to_remove._element)
                    new_row = table.add_row()
                    table.rows[row_to_remove_index]._element.addprevious(new_row._element)
                    self.read_text_widgets_and_write_file(new_row)
                doc.save(self.doc)
                self.table.clearContents()  # Чистим таблицу
                self.update_table_rows()
                self.read_file_table()      # Заполняем таблицу
                self.modific_doc_file()
            else:
                self.show_alert()
                self.alert_text.setText(f"Выберите колонку!")
                self.timer.setSingleShot(True)
                self.timer.timeout.connect(self.hide_alert)
                self.timer.start(3000)
        except PermissionError:
            self.show_alert()
            self.alert_text.setText(f"Файл был открыт!\nЗайкроте его и снова выполните сохранение данных в строке!")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)
        except Exception as ex:
            self.show_alert()
            self.alert_text.setText(f"Системная Ошибка! {ex}")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)

    # Функция, которая добавляет новую строку в таблицу
    def add_column_item(self):
        try:
            doc = Document(self.doc)
            table = doc.tables[self.table_id]
            if self.row:
                row_to_remove_index = self.row  # Нумерация строк начинается с 0
                print(row_to_remove_index)
                # Проход по всем таблицам в документе
                # Проверяем, что индекс удаляемой строки находится в пределах диапазона строк таблицы
                if 0 <= row_to_remove_index < len(table.rows):
                    new_row = table.add_row()
                    table.rows[row_to_remove_index]._element.addprevious(new_row._element)
                    self.read_text_widgets_and_write_file(new_row)


            else:
                new_row = table.add_row()
                self.read_text_widgets_and_write_file(new_row)

            # Сохранение изменений в документе
            doc.save(self.doc)
            self.table.clearContents()
            self.update_table_rows()
            self.read_file_table()
            self.modific_doc_file()
        except PermissionError:
            self.show_alert()
            self.alert_text.setText(f"Файл был открыт!\nЗайкроте его и снова выполните добавление строки!")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)
        except Exception as ex:
            self.show_alert()
            self.alert_text.setText(f"Системная Ошибка!\n{ex}")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)

    def add_merger_column(self):
        try:
            if self.row:
                doc = Document(self.doc)
                table = doc.tables[self.table_id]
                # Получаем строку по индексу
                row = table.rows[self.row]
                row_data = [cell.text.strip() for cell in row.cells]
                if len(set(row_data)) == 1:
                    # Получаем доступ к ячейке в указанной строке и столбце
                    cell = table.cell(self.row, 0)
                    # Устанавливаем новое значение содержимого ячейки
                    cell.text = self.lineEdit.text()
                    paragraph = cell.paragraphs[0]
                    # Выравниваем текст по центру
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Устанавливаем шрифт, стиль и размер шрифта для текста в объединенной ячейке
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    run.font.size = Pt(10)
                else:
                    # Индекс строки, перед которой нужно добавить новую строку (начиная с 0)
                    insert_row_index = self.row
                    # Проходим по всем таблицам в документе
                    for table in doc.tables:
                        # Проверяем, что индекс вставляемой строки находится в пределах диапазона строк таблицы
                        if 0 <= insert_row_index < len(table.rows):
                            # Создаем новую строку с 12 ячейками
                            new_row = table.add_row()
                            # Объединяем все ячейки новой строки
                            first_cell = new_row.cells[0]
                            last_cell = new_row.cells[-1]
                            merged_cell = first_cell.merge(last_cell)
                            # Вписываем значение в первую ячейку объединенного блока
                            merged_cell.text = self.lineEdit.text()
                            # Вставляем новую строку перед выбранной строкой
                            table.rows[insert_row_index]._element.addprevious(new_row._element)
                            # Получаем параграф в объединенной ячейке
                            paragraph = merged_cell.paragraphs[0]
                            # Выравниваем текст по центру
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            # Устанавливаем шрифт, стиль и размер шрифта для текста в объединенной ячейке
                            run = paragraph.runs[0]
                            run.font.name = 'Times New Roman'
                            run.font.bold = True
                            run.font.size = Pt(10)
            else:
                doc = Document(self.doc)
                # Получаем первую таблицу в документе
                table = doc.tables[self.table_id]
                # Добавляем новую строку в таблицу
                new_row = table.add_row()
                # Объединяем новые ячейки в одну
                merged_cell = table.cell(len(table.rows) - 1, 0).merge(
                    table.cell(len(table.rows) - 1, len(new_row.cells) - 1))
                new_row.cells[0].text = self.lineEdit.text()
                # Получаем параграф в объединенной ячейке
                paragraph = merged_cell.paragraphs[0]
                # Выравниваем текст по центру
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Устанавливаем шрифт, стиль и размер шрифта для текста в объединенной ячейке
                run = paragraph.runs[0]
                run.font.name = 'Times New Roman'
                run.font.bold = True
                run.font.size = Pt(10)
                # Сохраняем изменения
            doc.save(self.doc)
            self.update_table_rows()
            self.table.clearContents()
            self.read_file_table()
            self.modific_doc_file()
        except PermissionError:
            self.show_alert()
            self.alert_text.setText(f"Файл был открыт\nЗайкроте его и снова выполните добавление колонки!")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(3000)

    # Функция, которая удаляет строку в таблице
    def del_column(self):
        try:
            if self.row:
                # Загрузка документа
                doc = Document(self.doc)
                table = doc.tables[self.table_id]
                row_count = len(table.rows)
                # Индекс удаляемой строки (начиная с 0)
                row_to_remove_index = self.row  # Нумерация строк начинается с 0
                # Проход по всем таблицам в документе
                table = doc.tables[self.table_id]
                # Проверяем, что индекс удаляемой строки находится в пределах диапазона строк таблицы
                if 0 <= row_to_remove_index < len(table.rows):
                    # Удаляем строку из таблицы
                    parent = table.rows[row_to_remove_index]._element.getparent()
                    parent.remove(table.rows[row_to_remove_index]._element)

                doc.save(self.doc)
                self.table.clearContents()
                if row_count:
                    self.update_table_rows()
                self.read_file_table()
                self.modific_doc_file()
                self.row = 0
            else:
                self.show_alert()
                self.alert_text.setText(f"Выберите колонку!")
                self.timer.setSingleShot(True)
                self.timer.timeout.connect(self.hide_alert)
                self.timer.start(3000)
        except PermissionError:
            self.show_alert()
            self.alert_text.setText(f"Файл был открыт!\nЗайкроте его и снова выполните удаление строки!")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(3000)
        except Exception as ex:
            self.show_alert()
            self.alert_text.setText(f"Системная Ошибка!\n{ex}")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)

    # функция которая обновляет id или №п/п
    def update_table_rows(self):
        doc = Document(self.doc)
        idx = 1

        # Проходимся по всем таблицам в документе
        for table_id, table in enumerate(doc.tables):
            # Пропускаем таблицы, у которых не 12 ячеек
            if len(table.columns) != 12:
                continue
            elif table_id == 0:
                start_row = 3
            else:
                start_row = 1
            # Проходимся по строкам таблицы, начиная с четвертой строки
            for i, row in enumerate(table.rows):
                if i >= start_row:  # Начиная с четвертой строки
                    # Получаем значения всех ячеек текущей строки
                    cell_values = [cell.text for cell in row.cells]

                    # Проверяем, все ли значения ячеек одинаковы
                    if len(set(cell_values)) == 1:
                        # Пропускаем эту строку, так как все значения одинаковы
                        continue

                    # Получаем ячейку в первой колонке текущей строки
                    cell = row.cells[0]
                    # Очищаем содержимое ячейки
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.clear()
                    # Вставляем новый текст (индекс) с нужным форматированием
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(str(idx))
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    # Выравниваем текст по центру ячейки
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    idx += 1

        doc.save(self.doc)

    # Функция, которая блокирует некоторые виджеты, но оставляет режим термообработки
    def combining_heat_treatment(self):
        if self.main.user.posts == 'Технолог':
            if self.heat_treatment_mode_edit.toPlainText().strip():
                self.heating_medium_edit.setEnabled(False)
                self.annealing_temperature_edit.setEnabled(False)
                self.time_edit.setEnabled(False)
                self.cooling_conditions_edit.setEnabled(False)
                self.number_marshrut_edit.setEnabled(False)
                self.steps_deformac_edit.setEnabled(False)
            else:
                self.heating_medium_edit.setEnabled(True)
                self.annealing_temperature_edit.setEnabled(True)
                self.time_edit.setEnabled(True)
                self.cooling_conditions_edit.setEnabled(True)
                self.number_marshrut_edit.setEnabled(True)
                self.steps_deformac_edit.setEnabled(True)

    # Функция, которая блокирует вводить значения до и после диметра.
    def combining_diametr(self):
        if self.diametr_edit.text().strip():
            self.diametr_do_edit.setEnabled(False)
            self.diametr_posle_edit.setEnabled(False)
        else:
            self.diametr_do_edit.setEnabled(True)
            self.diametr_posle_edit.setEnabled(True)

    # Функция, которая блокирует комбинируемый диаметр
    def lock_combin(self):
        sender_widget = self.sender()
        if sender_widget is not None:
            if sender_widget.text().strip():
                self.diametr_edit.setEnabled(False)
            else:
                self.diametr_edit.setEnabled(True)

    # Функция, которая блокирует комбинируемый режим термообработки
    def lock_heart_combine(self):
        sender_widget = self.sender()
        if sender_widget is not None:
            if sender_widget.text().strip():
                self.heat_treatment_mode_edit.setEnabled(False)
            else:
                self.heat_treatment_mode_edit.setEnabled(True)

    # Функция которая чистит текст во всех виджетах
    def clear_widget(self):
        self.name_operation_box.setCurrentText(''), self.oborudovanie_edit.clear(), self.diametr_edit.clear()
        self.diametr_do_edit.clear(), self.diametr_posle_edit.clear(), self.number_marshrut_edit.clear()
        self.steps_deformac_edit.clear(), self.time_edit.clear()
        self.heat_treatment_mode_edit.clear(), self.heating_medium_edit.clear(), self.annealing_temperature_edit.clear()
        self.cooling_conditions_edit.clear(), self.note_edit.clear(), self.lineEdit.clear()
        self.row = 0
        self.cursor.setText(f"Курсор установлен на {self.row} строке")

    # Раскрывает список с таблицами
    def show_list_table(self):
        self.load_table_item()
        self.hide_table_window(True)

    # Раскрывает таблицу на максимум
    def unwrap_table(self):
        if self.open_scroll.isChecked():
            self.table.setGeometry(QtCore.QRect(20, 110, 1001, 701))
        else:
            self.table.setGeometry(QtCore.QRect(850, 110, 171, 701))

    # Функция, которая открывает или закрывает окно с информацией
    def open_info_window(self, bolean):
        self.info_window.setVisible(bolean)
        self.director.setVisible(bolean)
        self.director_edit.setVisible(bolean)
        self.number_card.setVisible(bolean)
        self.number_card_edit.setVisible(bolean)
        self.splav.setVisible(bolean)
        self.splav_edit.setVisible(bolean)
        self.gost.setVisible(bolean)
        self.gost_edit.setVisible(bolean)
        self.sto.setVisible(bolean)
        self.sto_edit.setVisible(bolean)
        self.vzamen.setVisible(bolean)
        self.vzamen_edit.setVisible(bolean)

    # Установка использования виджетов
    def widget_enable_true_or_false(self, bolean):
        self.open_list_table.setEnabled(bolean), self.open_scroll.setEnabled(bolean)
        self.reset_button.setEnabled(bolean), self.name_operation_box.setEnabled(bolean)
        self.oborudovanie_edit.setEnabled(bolean), self.diametr_edit.setEnabled(bolean)
        self.diametr_do_edit.setEnabled(bolean), self.diametr_posle_edit.setEnabled(bolean)
        self.number_marshrut_edit.setEnabled(bolean), self.steps_deformac_edit.setEnabled(bolean)
        self.heat_treatment_mode_edit.setEnabled(bolean), self.heating_medium_edit.setEnabled(bolean)
        self.annealing_temperature_edit.setEnabled(bolean), self.time_edit.setEnabled(bolean)
        self.cooling_conditions_edit.setEnabled(bolean), self.note_edit.setEnabled(bolean)
        self.lineEdit.setEnabled(bolean), self.button_column.setEnabled(bolean)
        self.button_save_column.setEnabled(bolean), self.button_add_column.setEnabled(bolean)
        self.button_del_column.setEnabled(bolean), self.save_file.setEnabled(bolean)
        self.open_file.setEnabled(bolean), self.delete_file.setEnabled(bolean)
        self.button_info_window.setEnabled(bolean)

    # Скрывает или показывает окно выбора таблица
    def hide_table_window(self, bolean):
        self.window.setVisible(bolean)
        self.table_window.setVisible(bolean)
        self.tables.setVisible(bolean)
        self.button_add_table.setVisible(bolean)
        self.button_del_table.setVisible(bolean)

    # Возвращает прежний вид виджетов
    def set_stylesheet_widget(self, stylesheet_green, stylesheet_red, open_window):
        self.save_file.setStyleSheet(stylesheet_green), self.open_file.setStyleSheet(stylesheet_green)
        self.delete_file.setStyleSheet(stylesheet_red), self.button_info_window.setStyleSheet(stylesheet_green)
        self.open_scroll.setStyleSheet("QPushButton {\n"
                                       "    background-color: white; \n"
                                       "    color: black;\n"
                                       "    border-radius: 10px;\n"
                                       "}\n"
                                       "\n"
                                       "QPushButton:hover {\n"
                                       "    background-color: green;\n"
                                       "}" if open_window else "QPushButton {\n"
                                       "    background-color: rgb(115, 115, 115); \n"
                                       "    color: black;\n"
                                       "    border-radius: 10px;\n"
                                       "}\n"
                                       "\n"
                                       "QPushButton:hover {\n"
                                       "    background-color: green;\n"
                                       "}"), self.button_column.setStyleSheet(stylesheet_green)
        self.button_save_column.setStyleSheet(stylesheet_green), self.button_add_column.setStyleSheet(stylesheet_green)
        self.button_del_column.setStyleSheet(stylesheet_green), self.open_list_table.setStyleSheet("QPushButton {\n"
                                            "    background-color: black; \n"
                                            "    color: white;\n"
                                            "    border-radius: 10px;\n"
                                            "}" if open_window else "QPushButton {\n"
                                       "    background-color: rgb(115, 115, 115); \n"
                                       "    color: black;\n"
                                       "    border-radius: 10px;\n"
                                       "}")
        self.reset_button.setStyleSheet("QPushButton {\n"
                                            "    background-color: black; \n"
                                            "    color: white;\n"
                                            "    border-radius: 10px;\n"
                                            "}" if open_window else "QPushButton {\n"
                                       "    background-color: rgb(115, 115, 115); \n"
                                       "    color: black;\n"
                                       "    border-radius: 10px;\n"
                                       "}")
        self.button_del_table.setStyleSheet("QPushButton {\n"
                                                    "    background-color: black; \n"
                                                    "    color: white;\n"
                                                    "    border-radius: 10px;\n"
                                                    "}\n"
                                                    "\n"
                                                    "QPushButton:hover {\n"
                                                    "    background-color: red;\n"
                                                    "}" if open_window else "QPushButton {\n"
                                                    "    background-color: rgb(115, 115, 115);; \n"
                                                    "    color: white;\n"
                                                    "    border-radius: 10px;\n"
                                                    "}\n"
                                                    "\n"
                                                    "QPushButton:hover {\n"
                                                    "    background-color: red;\n"
                                                    "}")