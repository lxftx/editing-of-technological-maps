import math
import os
import openpyxl


from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import QTimer
from PyQt5.QtWidgets import QWidget, QAbstractItemView, QFileDialog
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from openpyxl.styles import Font

from alert import AlertMessage
from docx import Document


class CalculationPage(QWidget, AlertMessage):
    def __init__(self, main):
        super().__init__()
        self.main = main
        self.calculation = self
        self.lml = []
        self.timer = QTimer()

        self.calculation_content = QtWidgets.QLabel(self.calculation)
        self.calculation_content.setGeometry(QtCore.QRect(10, 10, 1021, 861))
        font = QtGui.QFont()
        font.setPointSize(30)
        font.setBold(True)
        font.setWeight(75)
        self.calculation_content.setFont(font)
        self.calculation_content.setStyleSheet("background-color:rgb(240, 240, 240);\n"
                                               "border-radius: 30px;\n"
                                               "color: rgb(255, 255, 17);")
        self.calculation_content.setObjectName("calculation_content")
        self.how_load = QtWidgets.QLabel(self.calculation)
        self.how_load.setGeometry(QtCore.QRect(680, 20, 341, 141))
        self.how_load.setStyleSheet("background-color:rgb(240, 240, 240);\n"
                                    "border-radius: 30px;\n"
                                    "color: rgb(255, 255, 17);\n"
                                    "border: 1px solid black;")
        self.how_load.setObjectName("how_load")

        self.load_title = QtWidgets.QLabel(self.calculation)
        self.load_title.setGeometry(QtCore.QRect(690, 30, 321, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(14)
        self.load_title.setFont(font)
        self.load_title.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.load_title.setAlignment(QtCore.Qt.AlignCenter)
        self.load_title.setText("Выгрузка")
        self.load_title.setObjectName("load_title")

        self.word_label = QtWidgets.QLabel(self.calculation)
        self.word_label.setGeometry(QtCore.QRect(690, 70, 151, 20))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.word_label.setFont(font)
        self.word_label.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.word_label.setAlignment(QtCore.Qt.AlignCenter)
        self.word_label.setText("Word")
        self.word_label.setObjectName("word_label")

        self.excel_label = QtWidgets.QLabel(self.calculation)
        self.excel_label.setGeometry(QtCore.QRect(850, 70, 161, 21))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.excel_label.setFont(font)
        self.excel_label.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.excel_label.setAlignment(QtCore.Qt.AlignCenter)
        self.excel_label.setText("Excel")
        self.excel_label.setObjectName("excel_label")

        self.button_load_word = QtWidgets.QPushButton(self.calculation)
        self.button_load_word.setGeometry(QtCore.QRect(712, 110, 111, 31))
        self.button_load_word.setStyleSheet("border-radius: 10px;")
        icon1 = QtGui.QIcon()
        icon1.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "file-word_out.svg")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.button_load_word.setIcon(icon1)
        self.button_load_word.setIconSize(QtCore.QSize(24, 24))
        self.button_load_word.setObjectName("button_load_word")
        self.button_load_word.clicked.connect(self.load_word)

        self.button_load_excel = QtWidgets.QPushButton(self.calculation)
        self.button_load_excel.setGeometry(QtCore.QRect(880, 110, 111, 31))
        self.button_load_excel.setStyleSheet("border-radius: 10px;")
        icon2 = QtGui.QIcon()
        icon2.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "file-excel_out.svg")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.button_load_excel.setIcon(icon2)
        self.button_load_excel.setIconSize(QtCore.QSize(24, 24))
        self.button_load_excel.setObjectName("button_load_excel")
        self.button_load_excel.clicked.connect(self.load_excel)

        self.table_route = QtWidgets.QTableWidget(self.calculation)
        self.table_route.setGeometry(QtCore.QRect(30, 170, 981, 691))
        self.table_route.setObjectName("table_route")
        self.table_route.setColumnCount(3)
        self.table_route.setRowCount(0)
        self.table_route.setHorizontalHeaderLabels(["Проход", "Диаметр", "Деформация"])
        self.table_route.setEditTriggers(QAbstractItemView.NoEditTriggers)
        header = self.table_route.horizontalHeader()
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(20)
        header.setFont(font)
        header.setSectionResizeMode(QtWidgets.QHeaderView.Stretch)

        self.coef_deform_label = QtWidgets.QLabel(self.calculation)
        self.coef_deform_label.setGeometry(QtCore.QRect(40, 30, 211, 21))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.coef_deform_label.setFont(font)
        self.coef_deform_label.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.coef_deform_label.setText("Удлинение проволоки в -")
        self.coef_deform_label.setObjectName("coef_deform_label")

        self.coef_deform_edit = QtWidgets.QLineEdit(self.calculation)
        self.coef_deform_edit.setGeometry(QtCore.QRect(250, 25, 71, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.coef_deform_edit.setFont(font)
        self.coef_deform_edit.setObjectName("coef_deform_edit")
        self.coef_deform_edit.setPlaceholderText('%')

        self.width_label = QtWidgets.QLabel(self.calculation)
        self.width_label.setGeometry(QtCore.QRect(340, 30, 91, 21))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.width_label.setFont(font)
        self.width_label.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.width_label.setText("Диаметр -")
        self.width_label.setObjectName("width_label")

        self.width_edit = QtWidgets.QLineEdit(self.calculation)
        self.width_edit.setGeometry(QtCore.QRect(428, 25, 71, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.width_edit.setFont(font)
        self.width_edit.setObjectName("width_edit")

        self.count_label = QtWidgets.QLabel(self.calculation)
        self.count_label.setGeometry(QtCore.QRect(40, 70, 191, 21))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.count_label.setFont(font)
        self.count_label.setStyleSheet("background-color:rgb(240, 240, 240);")
        self.count_label.setText("Количество проходов - ")
        self.count_label.setObjectName("count_label")

        self.count_spinbox = QtWidgets.QSpinBox(self.calculation)
        self.count_spinbox.setMinimum(0)
        self.count_spinbox.setMaximum(100)
        self.count_spinbox.setGeometry(QtCore.QRect(232, 65, 71, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.count_spinbox.setFont(font)
        self.count_spinbox.setObjectName("count_spinbox")

        self.resert_items = QtWidgets.QPushButton(self.calculation)
        self.resert_items.setGeometry(QtCore.QRect(250, 120, 41, 31))
        self.resert_items.setStyleSheet("QPushButton {\n"
                                        "    background-color: black; \n"
                                        "    color: white;\n"
                                        "    border-radius: 10px;\n"
                                        "}")
        icon3 = QtGui.QIcon()
        icon3.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "reset_in.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.resert_items.setIcon(icon3)
        self.resert_items.setObjectName("resert_items")
        self.resert_items.clicked.connect(self.cleanItems)

        self.perform_calculation = QtWidgets.QPushButton(self.calculation)
        self.perform_calculation.setGeometry(QtCore.QRect(30, 122, 211, 31))
        font = QtGui.QFont()
        font.setFamily("Microsoft YaHei UI")
        font.setPointSize(12)
        self.perform_calculation.setFont(font)
        self.perform_calculation.setStyleSheet("QPushButton{\n"
                                               "    border-radius: 10px;\n"
                                               "}\n"
                                               "\n"
                                               "QPushButton:hover {\n"
                                               "    background-color: black;\n"
                                               "    color: white;\n"
                                               "}")
        self.perform_calculation.setText("Выполнить расчет")
        self.perform_calculation.setObjectName("perform_calculation")
        self.perform_calculation.clicked.connect(self.perform)

        self.alert = QtWidgets.QLabel(self.calculation)
        self.alert.setGeometry(QtCore.QRect(40, 270, 821, 201))
        self.alert.setStyleSheet("border-radius: 30px;")
        self.alert.setObjectName("alert_info")
        self.alert.setVisible(False)

        self.alert_title = QtWidgets.QLabel(self.calculation)
        self.alert_title.setGeometry(QtCore.QRect(40, 270, 961, 61))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(20)
        font.setBold(True)
        font.setWeight(75)
        self.alert_title.setFont(font)
        self.alert_title.setStyleSheet("border-top-left-radius: 30px;\n"
                                       "border-top-right-radius: 30px;\n"
                                       "background-color: rgb(227, 227, 227); \n"
                                       "color: rgb(179, 179, 179);")
        self.alert_title.setAlignment(QtCore.Qt.AlignCenter)
        self.alert_title.setObjectName("alert_info_title")
        self.alert_title.setText("Предупреждение")
        self.alert_title.setVisible(False)

        self.alert_text = QtWidgets.QLabel(self.calculation)
        self.alert_text.setGeometry(QtCore.QRect(40, 330, 961, 141))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(15)
        font.setUnderline(True)
        self.alert_text.setFont(font)
        self.alert_text.setStyleSheet("border-bottom-left-radius: 30px;\n"
                                      "border-bottom-right-radius: 30px;\n"
                                      "color: rgb(162, 162, 162);")
        self.alert_text.setAlignment(QtCore.Qt.AlignCenter)
        self.alert_text.setObjectName("alert_info_text")
        self.alert_text.setVisible(False)

        self.button_load_word.installEventFilter(self)
        self.button_load_excel.installEventFilter(self)
        self.resert_items.installEventFilter(self)

        if self.main.burger_button.isChecked():
            self.open_sidebar()

    def closeEvent(self, event):
        event.accept()

    def load_excel(self):
        try:
            if self.table_route.rowCount() * self.table_route.columnCount():
                file_path, _ = QFileDialog.getSaveFileName(None, "Сохранить документ Excel", "", "Документ Excel (*.xlsx)")
                if file_path:
                    workbook = openpyxl.Workbook()
                    sheet = workbook.active

                    # Добавляем тексты в первые две строки
                    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=5)
                    cell = sheet.cell(row=1, column=1)
                    cell.value = "Маршрут волочения"
                    cell.font = Font(name='Times New Roman', size=20, bold=True)

                    sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=5)
                    cell = sheet.cell(row=2, column=1)
                    cell.value = f"Удлинение проволоки в % {self.coef_deform_edit.text()}"
                    cell.font = Font(name='Times New Roman', size=14, bold=True)

                    row_count = self.table_route.rowCount()
                    column_count = self.table_route.columnCount()

                    # Записываем заголовки столбцов
                    for col in range(column_count):
                        header = self.table_route.horizontalHeaderItem(col)
                        sheet.cell(row=3, column=col + 1).value = header.text()
                        cell = sheet.cell(row=3, column=col + 1)
                        cell.value = header.text()
                        cell.font = Font(name='Times New Roman', size=12, bold=True)

                    # Записываем данные из таблицы
                    for row in range(row_count):
                        for col in range(column_count):
                            item = self.table_route.item(row, col)
                            if item is not None:
                                cell = sheet.cell(row=row + 4, column=col + 1)
                                cell.value = item.text()
                                cell.font = Font(name='Times New Roman', size=12)

                    workbook.save(file_path)
                    os.startfile(file_path)
            else:
                self.show_alert()
                self.alert_text.setText("Таблица пуста")
                self.timer.setSingleShot(True)
                self.timer.timeout.connect(self.hide_alert)
                self.timer.start(5000)
        except PermissionError:
            self.show_alert()
            self.alert_text.setText(f"Файл был открыт, зайкроте его и снова выполните загрузку в excel!")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)

    def add_header(self, text, font_size=20, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, doc=None):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        paragraph.alignment = alignment

    def load_word(self):
        try:
            if self.table_route.columnCount() * self.table_route.rowCount():
                file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить документ Word", "", "Документ Word (*.docx)")
                if file_path:
                    doc = Document()
                    self.add_header("Маршрут волочения", font_size=20, bold=True, doc=doc)
                    self.add_header("Удлинение проволоки в % 12.2", font_size=14, bold=True, doc=doc)
                    table = doc.add_table(rows=self.table_route.rowCount()+1, cols=self.table_route.columnCount())
                    table.style = "Table Grid"
                    headers = ["Проход", "Диаметр", "Деформация"]
                    for i, header in enumerate(headers):
                        cell = table.cell(0, i)
                        cell.text = header
                        paragraph = cell.paragraphs[0]
                        run = paragraph.runs[0]
                        run.bold = True
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for i, row in enumerate(self.lml, start=1):
                        for j, value in enumerate(row):
                            cell = table.cell(i, j)
                            cell.text = str(value)
                            paragraph = cell.paragraphs[0]
                            run = paragraph.runs[0]
                            run.font.name = "Times New Roman"  # Устанавливаем шрифт
                            run.font.size = Pt(12)  # Устанавливаем размер шрифта
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    doc.save(file_path)
                    os.startfile(file_path)
            else:
                self.show_alert()
                self.alert_text.setText("Таблица пуста")
                self.timer.setSingleShot(True)
                self.timer.timeout.connect(self.hide_alert)
                self.timer.start(5000)
        except PermissionError:
            self.show_alert()
            self.alert_text.setText(f"Файл был открыт, зайкроте его и снова выполните загрузку в word!")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)

    def get_path(self, dir, name):
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), dir, name)

    def perform(self):
        if self.coef_deform_edit.text():
            if self.width_edit.text():
                if self.count_spinbox.text():
                    l = [float(self.width_edit.text())]
                    lst = [l.append(l[x] * math.sqrt((float(self.coef_deform_edit.text()) + 100) / 100)) for x, _ in
                           enumerate(range(0, int(self.count_spinbox.text()) - 1))]
                    lml = [(len(l) - x, round(l[x],3),
                            "" if x + 2 > len(l) else round(((pow(l[x + 1], 2) - pow(l[x], 2)) / pow(l[x + 1], 2)), 3)) for x, _ in
                           enumerate(range(0, len(l)))]
                    self.table_route.setRowCount(len(lml))
                    for row_index, row_data in enumerate(lml):
                        for col_index, col_data in enumerate(row_data):
                            item = QtWidgets.QTableWidgetItem(str(col_data))
                            font = QtGui.QFont("Times New Roman", 12)
                            item.setFont(font)
                            item.setTextAlignment(QtCore.Qt.AlignCenter)
                            self.table_route.setItem(row_index, col_index, item)
                    self.lml = lml
                else:
                    self.show_alert()
                    self.alert_text.setText("Укажите количество проходов!")
                    self.timer.setSingleShot(True)
                    self.timer.timeout.connect(self.hide_alert)
                    self.timer.start(5000)
            else:
                self.show_alert()
                self.alert_text.setText("Укажите диаметр!")
                self.timer.setSingleShot(True)
                self.timer.timeout.connect(self.hide_alert)
                self.timer.start(5000)
        else:
            self.show_alert()
            self.alert_text.setText("Укажите удлинение проволоки!")
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.hide_alert)
            self.timer.start(5000)

    def cleanItems(self):
        self.table_route.clearContents()
        self.table_route.setRowCount(0)


    def eventFilter(self, watched, event):
        icon = QtGui.QIcon()
        if (event.type() == QtCore.QEvent.Enter) and (watched == self.button_load_word):
            icon.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "file-word_in.svg")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
            self.button_load_word.setIcon(icon)

        elif (event.type() == QtCore.QEvent.Leave) and (watched == self.button_load_word):
            icon.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "file-word_out.svg")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
            self.button_load_word.setIcon(icon)

        elif (event.type() == QtCore.QEvent.Enter) and (watched == self.button_load_excel):
            icon.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "file-excel_in.svg")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
            self.button_load_excel.setIcon(icon)

        elif (event.type() == QtCore.QEvent.Leave) and (watched == self.button_load_excel):
            icon.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "file-excel_out.svg")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
            self.button_load_excel.setIcon(icon)

        elif (event.type() == QtCore.QEvent.Enter) and (watched == self.resert_items):
            icon.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "reset_in.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
            self.resert_items.setIcon(icon)

        elif (event.type() == QtCore.QEvent.Leave) and (watched == self.resert_items):
            icon.addPixmap(QtGui.QPixmap(self.get_path(r'images\icon', "reset_out.png")), QtGui.QIcon.Normal, QtGui.QIcon.Off)
            self.resert_items.setIcon(icon)

        return super().eventFilter(watched, event)

    def open_sidebar(self):
        self.calculation_content.setGeometry(QtCore.QRect(10, 10, 881, 861))
        self.button_load_word.setGeometry(QtCore.QRect(572, 110, 111, 31))
        self.excel_label.setGeometry(QtCore.QRect(710, 70, 161, 21))
        self.how_load.setGeometry(QtCore.QRect(540, 20, 341, 141))
        self.load_title.setGeometry(QtCore.QRect(550, 30, 321, 31))
        self.button_load_excel.setGeometry(QtCore.QRect(740, 110, 111, 31))
        self.word_label.setGeometry(QtCore.QRect(550, 70, 151, 20))
        self.width_edit.setGeometry(QtCore.QRect(418, 20, 71, 31))
        self.coef_deform_label.setGeometry(QtCore.QRect(30, 25, 211, 21))
        self.count_spinbox.setGeometry(QtCore.QRect(222, 60, 71, 31))
        self.table_route.setGeometry(QtCore.QRect(30, 170, 841, 691))
        self.width_label.setGeometry(QtCore.QRect(330, 25, 91, 21))
        self.resert_items.setGeometry(QtCore.QRect(250, 115, 41, 31))
        self.count_label.setGeometry(QtCore.QRect(30, 65, 191, 21))
        self.perform_calculation.setGeometry(QtCore.QRect(30, 117, 211, 31))
        self.coef_deform_edit.setGeometry(QtCore.QRect(240, 20, 71, 31))
        self.alert.setGeometry(QtCore.QRect(40, 270, 821, 201))
        self.alert_title.setGeometry(QtCore.QRect(40, 270, 821, 61))
        self.alert_text.setGeometry(QtCore.QRect(40, 330, 821, 141))

    def close_sidebar(self):
        self.calculation_content.setGeometry(QtCore.QRect(10, 10, 1021, 861))
        self.button_load_word.setGeometry(QtCore.QRect(712, 110, 111, 31))
        self.excel_label.setGeometry(QtCore.QRect(850, 70, 161, 21))
        self.how_load.setGeometry(QtCore.QRect(680, 20, 341, 141))
        self.load_title.setGeometry(QtCore.QRect(690, 30, 321, 31))
        self.button_load_excel.setGeometry(QtCore.QRect(880, 110, 111, 31))
        self.word_label.setGeometry(QtCore.QRect(690, 70, 151, 20))
        self.table_route.setGeometry(QtCore.QRect(30, 170, 981, 691))
        self.coef_deform_label.setGeometry(QtCore.QRect(40, 30, 211, 21))
        self.coef_deform_edit.setGeometry(QtCore.QRect(250, 25, 71, 31))
        self.width_label.setGeometry(QtCore.QRect(340, 30, 91, 21))
        self.width_edit.setGeometry(QtCore.QRect(428, 25, 71, 31))
        self.count_label.setGeometry(QtCore.QRect(40, 70, 191, 21))
        self.count_spinbox.setGeometry(QtCore.QRect(232, 65, 71, 31))
        self.resert_items.setGeometry(QtCore.QRect(250, 120, 41, 31))
        self.perform_calculation.setGeometry(QtCore.QRect(30, 122, 211, 31))
        self.alert.setGeometry(QtCore.QRect(40, 270, 821, 201))
        self.alert_title.setGeometry(QtCore.QRect(40, 270, 961, 61))
        self.alert_text.setGeometry(QtCore.QRect(40, 330, 961, 141))